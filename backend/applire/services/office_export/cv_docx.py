# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""The direct ``python-docx`` CV writer (ADR-079, E057, US296).

``render_cv_docx`` is a **pure function**: no DB, no storage, no network —
``TailoredCVData`` plus a resolved accent colour and (optionally) resolved
photo bytes go in, ``.docx`` bytes come out. The caller (``services.cv``)
does all I/O before calling it, mirroring how ``get_cv_pdf`` composes
``get_cv_html`` — the export is rendered on demand from ``tailored_data`` and
no bytes are persisted (ADR-079 clause 8).

No HTML, no template engine, no subprocess anywhere in this module (ADR-079
clause 2): the section content is composed directly with the ``_common``
paragraph helpers.

**Section coverage is schema-driven, not a hand-written call sequence** —
at TWO levels:

1. ``render_cv_docx`` iterates ``TailoredCVData.model_fields`` and dispatches
   each TOP-LEVEL field to a registered renderer in ``_SECTION_RENDERERS``; a
   field with no registered renderer raises rather than being silently
   skipped.
2. ``_iter_leaf_paths`` walks the FULL nested field tree (into
   ``TailoredWorkEntry``, ``TailoredProjectEntry``, etc.), and
   ``test_every_tailored_cv_leaf_field_is_accounted_for`` requires every leaf
   to be in ``_RENDERED_LEAVES`` or ``_NOT_RENDERED_LEAVES`` (with a written
   reason). (1) alone is not enough: it is blind to fields nested inside a
   section's item model — ``work_history[].team_size`` /
   ``budget_managed`` / ``industry_context`` shipped unrendered past a green
   (1)-only suite for exactly that reason (all seven PDF templates render
   them; the .docx silently didn't — an undetected PDF/.docx content
   differential, against ADR-079's own premise that the export carries the
   same content set as the PDF). (2) is the SF-PROFILE.8 lesson applied here:
   the gate must cover the row's shape, not the fix's.
"""

import io
from typing import Callable

from docx.document import Document as DocxDocument
from docx.shared import Cm, RGBColor

from applire.schemas.cv import (
    TailoredCVData,
    TailoredCertification,
    TailoredEducationEntry,
    TailoredLanguage,
    TailoredProjectEntry,
    TailoredWorkEntry,
)
from applire.services.office_export._common import (
    _iter_leaf_paths,
    add_bullet,
    add_heading,
    add_paragraph,
    hex_to_rgb_color,
    new_document,
)
import logging

from applire.templates.filters import budget_display, education_title, month_year
from applire.templates.labels import cv_labels

logger = logging.getLogger(__name__)

PHOTO_WIDTH = Cm(3.2)

_DASH = " – "  # en dash, date ranges
_JOIN = " — "  # em dash, e.g. "Company — Role"


# ---------------------------------------------------------------------------
# `_iter_leaf_paths` (the nested-field schema walker behind the coverage
# guard below) moved to `_common.py` in E057 task 1.4 (US297, ADR-066) once
# the letter writer needed the identical walk over `LetterData` — it was
# already document-kind agnostic, so this is a pure relocation. Imported
# above; re-used here unchanged.
# ---------------------------------------------------------------------------


def _has_text(*values: str | None) -> bool:
    """True if at least one of `values` is non-blank after stripping."""
    return any(v and v.strip() for v in values)


def _join_nonblank(parts: list[str | None], sep: str = _JOIN) -> str:
    return sep.join(p.strip() for p in parts if p and p.strip())


def _format_date_range(start: str | None, end: str | None, labels: dict) -> str:
    """'01/2020 – 06/2022', '07/2022 – heute'/'Present' (end is None but start
    is set), or '' when both are blank. Never renders the literal 'None'.

    Formatted with the SAME ``month_year`` filter all seven PDF templates use
    (`{{ job.start_date | month_year }}`), not a second rule of its own. The
    export previously emitted the storage format (`2017-04`), so one document
    showed two different dates depending on which file the reader opened —
    SF-EXPORT.2's failure mode, and two implementations of one display rule
    (ADR-066). Every containment assertion in this suite stayed green through
    it, because the raw value IS present either way: the tests could not see
    formatting at all. Caught by converting a real export and looking at it.
    """
    start = month_year((start or "").strip())
    end = month_year((end or "").strip())
    if not start and not end:
        return ""
    if not end:
        end = labels["present"]
    if start and end:
        return f"{start}{_DASH}{end}"
    return start or end


def _role_facts_line(entry: TailoredWorkEntry, labels: dict, lang: str) -> str:
    """#328 (ADR-062 clause 1) per-role quantified facts, rendered as
    deterministic document furniture — 'Label: value' pairs, never composed
    into a sentence, matching how all seven PDF templates render them via
    the same `role_team_size` / `role_budget` / `role_industry` labels.

    `team_size` is guarded on `is not None`, never truthiness: it is
    `int | None` and 0 is a stated, meaningful fact — "None means 'not
    stated' — 0 is a valid team_size" (schemas/cv.py:140). `if
    entry.team_size:` would silently drop a real zero.
    """
    parts = []
    if entry.team_size is not None:
        parts.append(f"{labels['role_team_size']}: {entry.team_size}")
    # Through the SAME filter every template applies
    # (`job.budget_managed | budget_display(lang)`), and — like the templates'
    # `{% if budget_text %}` — the line is omitted when the filter rejects the
    # value. `budget_display("6000000")` returns "" precisely because a bare
    # unit-less number shipped as #382; printing it raw here would reintroduce
    # that defect on the export.
    budget_text = budget_display(entry.budget_managed, lang)
    if _has_text(budget_text):
        parts.append(f"{labels['role_budget']}: {budget_text.strip()}")
    if _has_text(entry.industry_context):
        parts.append(f"{labels['role_industry']}: {entry.industry_context.strip()}")
    return "   |   ".join(parts)


def _project_has_content(project: TailoredProjectEntry) -> bool:
    return _has_text(project.name) or any(_has_text(b) for b in project.bullets)


def _work_entry_has_content(entry: TailoredWorkEntry) -> bool:
    if _has_text(entry.company, entry.role, entry.start_date, entry.end_date):
        return True
    # See _role_facts_line: 0 is a stated team_size, not "nothing to show".
    if entry.team_size is not None:
        return True
    if _has_text(entry.budget_managed, entry.industry_context):
        return True
    if any(_has_text(b) for b in entry.bullets):
        return True
    return any(_project_has_content(p) for p in entry.projects)


def _education_has_content(entry: TailoredEducationEntry) -> bool:
    return _has_text(entry.institution, entry.degree, entry.field, entry.start_date, entry.end_date)


def _language_has_content(entry: TailoredLanguage) -> bool:
    return _has_text(entry.language, entry.level)


def _certification_has_content(cert: TailoredCertification) -> bool:
    return _has_text(cert.name, cert.issuing_organization, cert.date_obtained, cert.expiry_date)


def _render_project(document: DocxDocument, project: TailoredProjectEntry) -> None:
    add_paragraph(document, project.name, bold=True)
    for bullet in project.bullets:
        add_bullet(document, bullet)


# ---------------------------------------------------------------------------
# Section renderers — one per TailoredCVData field. Every renderer takes the
# same five arguments so the dispatch loop below stays uniform; renderers
# that don't need `photo_bytes` simply ignore it.
# ---------------------------------------------------------------------------


def _render_contact(document, tailored, labels, color, photo_bytes, lang) -> None:
    contact = tailored.contact
    if photo_bytes:
        try:
            document.add_picture(io.BytesIO(photo_bytes), width=PHOTO_WIDTH)
        except Exception:
            # python-docx reads only bmp/gif/jpeg/png/tiff. `services/cv.py`'s
            # _PHOTO_MIME declares **webp** a supported stored-photo type, and
            # Chromium renders it natively on the PDF path — so a candidate
            # with a .webp photo got a correct PDF and an HTTP 500 on every
            # .docx download (routers/cv.py maps any raise to 500). Corrupt or
            # truncated bytes fail the same way. Degrade exactly as the PDF
            # path already does for a missing file: omit the photo, keep the
            # document. Never let an image cost the candidate their text.
            logger.warning("office export: photo omitted, unreadable by python-docx")
    add_heading(document, contact.name, 1, color)
    details = _join_nonblank(
        [contact.location, contact.phone, contact.email, contact.linkedin],
        sep="   |   ",
    )
    add_paragraph(document, details)


def _render_summary(document, tailored, labels, color, photo_bytes, lang) -> None:
    if not _has_text(tailored.summary):
        return
    add_heading(document, labels["summary"], 2, color)
    add_paragraph(document, tailored.summary)


def _render_work_history(document, tailored, labels, color, photo_bytes, lang) -> None:
    if not any(_work_entry_has_content(e) for e in tailored.work_history):
        return
    add_heading(document, labels["experience"], 2, color)
    for entry in tailored.work_history:
        if not _work_entry_has_content(entry):
            continue
        header = _join_nonblank([entry.company, entry.role])
        add_heading(document, header, 3, color)
        add_paragraph(document, _format_date_range(entry.start_date, entry.end_date, labels), italic=True)
        add_paragraph(document, _role_facts_line(entry, labels, lang), italic=True)
        for bullet in entry.bullets:
            add_bullet(document, bullet)
        for project in entry.projects:
            _render_project(document, project)


def _render_skills(document, tailored, labels, color, photo_bytes, lang) -> None:
    if not any(_has_text(s) for s in tailored.skills):
        return
    add_heading(document, labels["skills"], 2, color)
    for skill in tailored.skills:
        add_bullet(document, skill)


def _render_education(document, tailored, labels, color, photo_bytes, lang) -> None:
    if not any(_education_has_content(e) for e in tailored.education):
        return
    add_heading(document, labels["education"], 2, color)
    for entry in tailored.education:
        if not _education_has_content(entry):
            continue
        # `education_title` dedupes a degree that already names its field —
        # "Industriemeister Metall" + field "Metall" renders once, not twice.
        # The filter exists because #548 shipped exactly that redundancy and a
        # blind reviewer flagged it on a real run; every CV template applies it
        # (`{{ edu.degree | education_title(edu.field) }}`), so the export must
        # too or it reintroduces a closed defect on a new surface.
        header = _join_nonblank([entry.institution, education_title(entry.degree, entry.field)])
        add_paragraph(document, header, bold=True)
        add_paragraph(document, _format_date_range(entry.start_date, entry.end_date, labels), italic=True)


def _render_languages(document, tailored, labels, color, photo_bytes, lang) -> None:
    if not any(_language_has_content(e) for e in tailored.languages):
        return
    add_heading(document, labels["languages"], 2, color)
    for entry in tailored.languages:
        add_bullet(document, _join_nonblank([entry.language, entry.level]))


def _render_projects(document, tailored, labels, color, photo_bytes, lang) -> None:
    if not any(_project_has_content(p) for p in tailored.projects):
        return
    add_heading(document, labels["projects"], 2, color)
    for project in tailored.projects:
        _render_project(document, project)


def _render_certifications(document, tailored, labels, color, photo_bytes, lang) -> None:
    if not any(_certification_has_content(c) for c in tailored.certifications):
        return
    add_heading(document, labels["certifications"], 2, color)
    for cert in tailored.certifications:
        # Same shared filter as every other date on this document.
        dates = _join_nonblank(
            [month_year(cert.date_obtained), month_year(cert.expiry_date)], sep=_DASH
        )
        add_bullet(document, _join_nonblank([cert.name, cert.issuing_organization, dates]))


_SECTION_RENDERERS: dict[str, Callable] = {
    "contact": _render_contact,
    "summary": _render_summary,
    "work_history": _render_work_history,
    "skills": _render_skills,
    "education": _render_education,
    "languages": _render_languages,
    "projects": _render_projects,
    "certifications": _render_certifications,
}

# `show_photo` is a boolean modifier of `contact` (whether the caller-resolved
# photo is embedded there), never a section of its own — the caller already
# gates `photo_bytes` on `show_photo` before calling render_cv_docx.
_NON_SECTION_FIELDS: frozenset[str] = frozenset({"show_photo"})


# ---------------------------------------------------------------------------
# Nested-leaf coverage registries. Every leaf `_iter_leaf_paths` finds in
# TailoredCVData's full field tree must be in EXACTLY ONE of these two sets
# (test_every_tailored_cv_leaf_field_is_accounted_for enforces it) — a leaf
# in neither is a silent omission; "structural, not content" is a real
# category (e.g. `show_photo`) but it must be a WRITTEN decision, never an
# absence.
# ---------------------------------------------------------------------------

_RENDERED_LEAVES: frozenset[str] = frozenset({
    "contact.name",
    "contact.email",
    "contact.phone",
    "contact.location",
    "contact.linkedin",
    "summary",
    "work_history[].company",
    "work_history[].role",
    "work_history[].start_date",
    "work_history[].end_date",
    "work_history[].bullets",
    "work_history[].team_size",
    "work_history[].budget_managed",
    "work_history[].industry_context",
    "work_history[].projects[].name",
    "work_history[].projects[].bullets",
    "skills",
    "education[].institution",
    "education[].degree",
    "education[].field",
    "education[].start_date",
    "education[].end_date",
    "languages[].language",
    "languages[].level",
    "projects[].name",
    "projects[].bullets",
    "certifications[].name",
    "certifications[].issuing_organization",
    "certifications[].date_obtained",
    "certifications[].expiry_date",
})

_NOT_RENDERED_LEAVES: dict[str, str] = {
    "contact.photo_url": (
        "Not read by this pure function at all: the caller "
        "(services.cv.get_cv_docx) resolves it via the existing "
        "_resolve_photo_data_uri storage lookup BEFORE calling "
        "render_cv_docx, and passes the result as the separate "
        "`photo_bytes` parameter. The photo IS embedded (_render_contact, "
        "when photo_bytes is truthy) — just never by this field directly."
    ),
    "work_history[].id": (
        "Internal correlation key carried from the source WorkEntry.id so "
        "services.cv._nest_projects can match a tailored entry back to its "
        "source (schemas/cv.py) — a plumbing id, never user-facing content "
        "on the candidate's own CV."
    ),
    "show_photo": (
        "A boolean modifier of contact.photo_url (whether the caller "
        "resolves and passes photo_bytes at all), not a content field of "
        "its own — consumed one layer up in services.cv.get_cv_docx, "
        "mirrored by _NON_SECTION_FIELDS above for the same reason."
    ),
}


def render_cv_docx(
    tailored: TailoredCVData,
    *,
    lang: str,
    accent_color: str,
    photo_bytes: bytes | None = None,
) -> bytes:
    """Render `tailored` as a `.docx` file, returned as bytes.

    Pure function — no I/O. `photo_bytes` is the already-resolved photo (or
    None if absent/unreadable/not shown); the caller decides whether to
    resolve and pass it based on `tailored.show_photo`.
    """
    labels = cv_labels(lang)
    document = new_document()
    color: RGBColor = hex_to_rgb_color(accent_color)

    for field_name in TailoredCVData.model_fields:
        if field_name in _NON_SECTION_FIELDS:
            continue
        renderer = _SECTION_RENDERERS.get(field_name)
        if renderer is None:
            raise RuntimeError(
                f"office_export.cv_docx: TailoredCVData field {field_name!r} has no "
                "registered renderer in _SECTION_RENDERERS — a schema field would "
                "silently go unexported from the .docx writer."
            )
        renderer(document, tailored, labels, color, photo_bytes, lang)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
