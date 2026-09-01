# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""E057 Task 1.1 (US296, #637, ADR-079) — the .docx text extractor and the audit seam.

``extract_docx_text`` walks a ``.docx``'s ``word/document.xml`` body in DOCUMENT
ORDER — the same physical order the XML lists elements in, regardless of how
deeply a run is nested (a top-level paragraph, a table cell's paragraph, an
anchored text box's paragraph are all just descendants encountered where they
occur). This is deliberately NOT ``document.paragraphs`` then ``document.tables``
concatenated: python-docx's convenience accessors only enumerate paragraph/table
BLOCK ITEMS directly under the body (or a cell) — a table nested between two
paragraphs reports as if it came LAST, and an anchored/inline text box's content
is invisible to either accessor (the drawing that carries it is neither a
``<w:p>`` nor a ``<w:tbl>``; a paragraph that only contains a drawing simply has
an empty ``.text``). Measured in
``tests/unit/test_office_export_extract.py::test_naive_paragraphs_reader_loses_table_content``
and ``::test_naive_paragraphs_reader_loses_text_box_content`` — this is the exact
naive shape ``cv_parser.extract_text_from_docx`` (the CV *import* path) already
uses; see this task's report for the measured cost of that on a tabular CV.

Joining rule, and why (US296 task spec's own framing: "paragraph -> newline, run
-> no separator, cell -> ?"):

* paragraph (``</w:p>`` closing) -> one ``"\\n"`` appended. A table cell's or a
  text box's LAST paragraph closes exactly like any other paragraph, so a cell
  or text-box boundary is already newline-terminated by construction (every
  ``<w:tc>``/``<w:txbxContent>`` contains >= 1 ``<w:p>`` by the OOXML schema) —
  no separate "cell separator" rule is needed or added.
* ``<w:t>`` run text -> appended with NO separator, ever, between runs or
  between a run and the next. Word routinely splits one visible WORD across
  multiple runs with no real space between them at all (spell-check proofing
  marks, mid-word bold/italic, a saved revision) — inserting a separator
  between runs would silently corrupt exactly the words ``_find`` most needs
  intact. See ``test_split_run_produces_no_spurious_separator`` and
  measurement (a) below.
* ``<w:tab/>`` -> one literal space. Unlike a run split (a formatting/proofing
  ARTEFACT with no separator intended), a tab is an author-placed gap — Word
  resume templates commonly right-align a date via a tab stop rather than a
  table cell.
* ``<w:br/>`` / ``<w:cr/>`` (a soft line break, not a new paragraph) -> ``"\\n"``,
  matching how it renders.
* an ``<mc:Fallback>`` subtree -> skipped entirely, not recursed into. Word
  wraps a text box in ``<mc:AlternateContent>`` with a modern DrawingML
  ``<mc:Choice Requires="wps">`` branch AND a legacy VML ``<mc:Fallback>``
  branch carrying the SAME text (for readers that don't understand the modern
  branch) — walking both would report every text box's words twice. Measured
  in ``test_alternate_content_text_box_does_not_duplicate_text``.

Why this is safe for ``_find`` (US296 measurement (a), ADR-079 clause 4, the
2026-08-31 boundary re-baseline note): ``_find``'s #399 kerning tolerance only
ever ADDS forgiveness for extra ASCII spaces BETWEEN the needle's own
characters — it can never skip a needle character that is genuinely absent from
the haystack. So even on the rare occasion the loose fallback fires, it cannot
turn a genuinely shortened/altered bullet (the #634 shape) into a false match.
Measured, not assumed — see ``tests/unit/test_office_export_extract.py``'s
``test_measurement_a_*`` and ``test_measurement_b_*`` functions, and this
task's report for the numbers.

── The audit seam ──────────────────────────────────────────────────────────────
``audit_cv_docx`` / ``audit_cover_letter_docx`` extract the produced ``.docx``'s
text and feed it to the UNCHANGED ``_audit_cv_text`` / ``_audit_letter_text``
(🔒 Architecture Boundary, ADR-066 / ADR-079 clause 4) — same predicate, same
checks, over ``.docx``-derived text instead of PDF-derived text. Every other
keyword argument (``ledger``, ``vault_text_norm``, ``vault_skill_forms``,
``pins``, ``truth_floor_hits``) passes straight through, so this is a full-
fidelity feed, not a stripped-down one.

*** RESOLVED 2026-09-01 (PO) — this seam reports the band explicitly ***
ADR-079 clause 4 requires the page-length band to be reported ``not_applicable``
WITH its reason, in its own bucket — never folded into an ``X of Y`` rollup, and
never silently omitted. When this module was first written that was an OPEN
QUESTION: ``ATSCheck.status`` was ``Literal["pass", "fail"]``, and passing no
``page_count`` — the only move open to a caller forbidden from editing the two
🔒-protected audit functions — makes the page-length block SKIP entirely. The
check is then not reported as anything; it is simply ABSENT from ``checks``,
indistinguishable in shape from a report persisted before the band existed, and
invisible to ``passed`` and ``failed`` alike. That is the SAME failure class as
#634: an instrument's silence read as evidence about something it never examined.

Two ways out were costed: (a) give the two audit functions a defaulted "this
band does not apply" parameter — a boundary exception, but keeps "what a
page-length row looks like" decided in exactly ONE place (ADR-066), which is
what the boundary existed to protect; or (b) construct the check here, which
respects the boundary's letter but opens a THIRD ``page-length`` construction
site to hand-sync, and fixes only this seam — any future ``page_count``-less
caller would reproduce the identical silent absence with nothing to stop it.

**The founder chose (a).** ``_audit_cv_text`` / ``_audit_letter_text`` now take
``page_band_not_applicable``, defaulting to today's behaviour for all ~90
existing callers, and both seams below set it. ``ATSCheck.status`` gained
``not_applicable`` as a third state, counted in its own ``ATSReport``
bucket and honoured by every consumer down to the frontend panel.

"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from lxml.etree import QName

from applire.schemas.ats import ATSReport
from applire.schemas.cv import TailoredCVData
from applire.services.ats_audit import _audit_cv_text, _audit_letter_text


def _local_name(tag: str) -> str:
    """Namespace-agnostic local tag name, e.g. ``'{...}t'`` -> ``'t'``."""
    return QName(tag).localname


def _walk_docx_body(element: Any, buf: list[str]) -> None:
    """Depth-first, document-order walk of a docx XML element, appending text
    fragments to ``buf`` per the joining rule documented at module level.

    Generic on purpose: table cells, text-box paragraphs and top-level
    paragraphs are all just ``<w:p>`` descendants encountered wherever they
    physically occur — no special-casing per container is needed, which is
    exactly what keeps document order correct without reconstructing it from
    ``.paragraphs`` + ``.tables``.
    """
    tag = _local_name(element.tag)
    if tag == "Fallback":
        # mc:AlternateContent's legacy-VML branch — the mc:Choice branch
        # (walked normally, no special case needed) already carries this
        # text; recursing here would report it twice.
        return
    if tag == "t":
        buf.append(element.text or "")
    elif tag == "tab":
        buf.append(" ")
    elif tag in ("br", "cr"):
        buf.append("\n")
    for child in element:
        _walk_docx_body(child, buf)
    if tag == "p":
        buf.append("\n")


def extract_docx_text(data: bytes) -> str:
    """All ``w:t`` run text in a ``.docx``'s body, in document order, including
    table cells and text boxes (see module docstring for the joining rule and
    why the naive ``document.paragraphs`` reader misses both).
    """
    document = Document(BytesIO(data))
    buf: list[str] = []
    _walk_docx_body(document.element.body, buf)
    return "".join(buf)


def audit_cv_docx(
    docx_bytes: bytes,
    tailored: TailoredCVData,
    keywords: list[str],
    ledger: list[dict[str, Any]] | None = None,
    vault_text_norm: str | None = None,
    vault_skill_forms: list[str] | None = None,
    pins: list | None = None,
) -> ATSReport:
    """Audit a produced CV ``.docx`` against the structured CV data and keywords —
    the ``.docx`` twin of ``ats_audit.audit_cv`` (which audits a PDF).

    Extracts via :func:`extract_docx_text` and feeds ``_audit_cv_text`` with
    ``page_band_not_applicable=True``: a ``.docx`` has no pages until a word
    processor lays it out, so the ADR-051 band is genuinely inapplicable and is
    reported as such rather than omitted (ADR-079 cl. 4). Every other keyword
    argument mirrors :func:`applire.services.ats_audit.audit_cv`.
    """
    text = extract_docx_text(docx_bytes)
    return _audit_cv_text(
        text,
        tailored,
        keywords,
        ledger,
        page_band_not_applicable=True,
        vault_text_norm=vault_text_norm,
        vault_skill_forms=vault_skill_forms,
        pins=pins,
    )


def audit_cover_letter_docx(
    docx_bytes: bytes,
    letter_data: dict[str, Any],
    keywords: list[str],
    ledger: list[dict[str, Any]] | None = None,
    vault_text_norm: str | None = None,
    pins: list | None = None,
    truth_floor_hits: set[str] | frozenset[str] = frozenset(),
) -> ATSReport:
    """Audit a produced cover-letter ``.docx`` against the structured letter data
    and keywords — the ``.docx`` twin of ``ats_audit.audit_cover_letter``.

    As in :func:`audit_cv_docx`, the page-length band is reported
    ``not_applicable`` rather than omitted (ADR-079 cl. 4).
    """
    text = extract_docx_text(docx_bytes)
    return _audit_letter_text(
        text,
        letter_data,
        keywords,
        ledger,
        page_band_not_applicable=True,
        vault_text_norm=vault_text_norm,
        pins=pins,
        truth_floor_hits=truth_floor_hits,
    )
