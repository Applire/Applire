# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""The direct ``python-docx`` cover-letter writer (ADR-079, E057, US297).

``render_letter_docx`` is a **pure function**: no DB, no storage, no network —
``LetterData`` plus a resolved accent colour go in, ``.docx`` bytes come out.
The caller (``services.cover_letter``) does all I/O before calling it,
mirroring how ``render_cv_docx`` is called from ``services.cv.get_cv_docx``
(US296) — rendered on demand from ``letter_data``, no bytes persisted
(ADR-079 clause 8).

No HTML, no template engine, no subprocess anywhere in this module (ADR-079
clause 2): section content is composed directly with the ``_common``
paragraph helpers — the same ones the CV writer uses, imported unchanged.

**Section coverage is schema-driven, not a hand-written call sequence** — at
the SAME two levels ``cv_docx.py`` uses:

1. ``render_letter_docx`` iterates ``LetterData.model_fields`` and dispatches
   each TOP-LEVEL field (``header``, ``recipient``, ``body``, ``signature``)
   to a registered renderer in ``_SECTION_RENDERERS``; a field with no
   registered renderer raises rather than being silently skipped.
2. ``_iter_leaf_paths`` (``_common.py`` — moved there from ``cv_docx.py`` in
   this same task, once this second writer needed the identical walk)
   walks the FULL nested field tree, and
   ``test_every_letterdata_leaf_field_is_accounted_for`` requires every leaf
   to be in ``_RENDERED_LEAVES`` or ``_NOT_RENDERED_LEAVES`` (with a written
   reason).

**Two deliberate asymmetries with the CV writer — both grounded in the
schema and the seven existing letter templates, not invented here:**

* **No ``photo_bytes`` parameter.** ``LetterHeader.photo_url`` carries the
  schema's own docstring (``schemas/cover_letter.py``): "Present in the
  writer shape but never rendered by any letter template. The
  render_document entry point STRIPS it (storage-read safety, US250)."
  Grepping all seven ``*_letter.html.j2`` templates confirms it: not one of
  them renders ``letter.header.photo_url``. Unlike the CV (where the photo
  IS content, just resolved one layer up by the caller), there is no letter
  content this leaf could ever contribute — so this writer takes no photo
  parameter at all, and ``header.photo_url`` is this module's one
  ``_NOT_RENDERED_LEAVES`` entry.
* **One heading, not three.** The CV has labelled section headings sourced
  from ``cv_labels()`` (``experience``, ``skills``, ...). ``cover_letter_
  labels()`` has no equivalent key naming a "recipient"/"body"/"signature"
  section, and no letter template renders a label for those blocks either
  (checked against all seven — the only per-block labels that exist at all,
  ``labels.email``/``.phone``/``.address``, are ``creative_sidebar``'s own
  sidebar chrome, not used by the other six templates, and — matching how
  ``cv_docx._render_contact`` already renders CV contact details unlabelled —
  not reproduced here either). Inventing hardcoded heading text would
  violate the i18n contract (never hard-code DE/EN strings). So this writer
  applies the SAME accent-colour-on-headings RULE as the CV (``add_heading``,
  a character-run colour, never a border or shaded box) to the one heading a
  letter naturally has: the sender's own name (``header.name``, Heading 1).
"""

import io
from typing import Callable

from docx.document import Document as DocxDocument
from docx.shared import RGBColor

from applire.schemas.cover_letter import LetterData
from applire.services.office_export._common import (
    add_heading,
    add_paragraph,
    hex_to_rgb_color,
    new_document,
)
from applire.templates.labels import cover_letter_labels

# `_iter_leaf_paths` is not imported here: this module's own code never
# calls it (only the coverage-guard tests do, importing it straight from
# `_common` — the walker's real home since this task's move, ADR-066).


def _closing_line(closing: str | None, labels: dict) -> str:
    """``signature.closing`` + the language's closing punctuation
    (``cover_letter_labels()['closing_punctuation']``), concatenated with NO
    separator — matching every letter template's own
    ``{{ letter.signature.closing }}{{ labels.closing_punctuation }}``
    (#307: German takes no comma after the Grussformel; English does). Blank
    or ``None`` closing returns ``""`` so the caller's ``add_paragraph``
    no-ops rather than rendering a lone punctuation mark on its own line.
    """
    if not closing or not closing.strip():
        return ""
    return f"{closing.strip()}{labels['closing_punctuation']}"


# ---------------------------------------------------------------------------
# Section renderers — one per LetterData field. Every renderer takes the same
# four arguments so the dispatch loop below stays uniform (no photo_bytes —
# see module docstring for why the letter writer, unlike the CV's, takes
# none at all).
# ---------------------------------------------------------------------------


def _render_header(document: DocxDocument, letter: LetterData, labels: dict, color: RGBColor) -> None:
    header = letter.header
    add_heading(document, header.name, 1, color)
    add_paragraph(document, header.address)
    add_paragraph(document, header.phone)
    add_paragraph(document, header.email)


def _render_recipient(document: DocxDocument, letter: LetterData, labels: dict, color: RGBColor) -> None:
    recipient = letter.recipient
    add_paragraph(document, recipient.name)
    add_paragraph(document, recipient.title)
    add_paragraph(document, recipient.company)
    add_paragraph(document, recipient.address)
    add_paragraph(document, recipient.date)


def _render_body(document: DocxDocument, letter: LetterData, labels: dict, color: RGBColor) -> None:
    for paragraph in letter.body.paragraphs:
        add_paragraph(document, paragraph)


def _render_signature(document: DocxDocument, letter: LetterData, labels: dict, color: RGBColor) -> None:
    signature = letter.signature
    add_paragraph(document, _closing_line(signature.closing, labels))
    add_paragraph(document, signature.name)


_SECTION_RENDERERS: dict[str, Callable] = {
    "header": _render_header,
    "recipient": _render_recipient,
    "body": _render_body,
    "signature": _render_signature,
}

# Every LetterData top-level field IS a real section — unlike TailoredCVData's
# `show_photo` boolean modifier, LetterData has no non-section field. Kept
# (empty) for structural symmetry with cv_docx.py's coverage-guard pattern,
# so a future non-section field added to LetterData has a documented place
# to go rather than forcing a new pattern into render_letter_docx's loop.
_NON_SECTION_FIELDS: frozenset[str] = frozenset()


# ---------------------------------------------------------------------------
# Nested-leaf coverage registry. Every leaf `_iter_leaf_paths` finds in
# LetterData's full field tree must be in EXACTLY ONE of these two sets
# (test_every_letterdata_leaf_field_is_accounted_for enforces it) — a leaf
# in neither is a silent omission; a leaf in both is a self-contradiction.
# ---------------------------------------------------------------------------

_RENDERED_LEAVES: frozenset[str] = frozenset({
    "header.name",
    "header.address",
    "header.phone",
    "header.email",
    "recipient.name",
    "recipient.title",
    "recipient.company",
    "recipient.address",
    "recipient.date",
    "body.paragraphs",
    "signature.closing",
    "signature.name",
})

_NOT_RENDERED_LEAVES: dict[str, str] = {
    "header.photo_url": (
        "schemas/cover_letter.py's own docstring: present in the writer "
        "shape but never rendered by any letter template; render_document "
        "STRIPS it before an agent-authored letter is even persisted "
        "(storage-read safety, US250). Confirmed against all seven "
        "*_letter.html.j2 templates by grep — none references it. Unlike "
        "the CV, this writer takes no photo_bytes parameter at all: there "
        "is no content this leaf could ever contribute to a letter .docx."
    ),
}


def _document_title(labels: dict, letter: LetterData) -> str:
    """``"Bewerbung – Catherine O'Brien"`` / ``"Application – …"`` (ADR-085 cl. 5).

    The letter templates carry no ``<title>`` at all — an asymmetry with the CV
    side that is recorded as a collector line, not fixed here — so this uses the
    same language-keyed ``subject_prefix`` the letter's own subject line uses,
    which keeps the `.docx` title language-correct without inventing chrome.
    """
    name = (getattr(letter.signature, "name", None) or "").strip()
    kind = labels["subject_prefix"]
    return f"{kind} – {name}" if name else kind


def render_letter_docx(letter: LetterData, *, lang: str, accent_color: str) -> bytes:
    """Render `letter` as a `.docx` file, returned as bytes.

    Pure function — no I/O. See module docstring for the two deliberate
    asymmetries with `render_cv_docx`.
    """
    labels = cover_letter_labels(lang)
    document = new_document(title=_document_title(labels, letter))
    color: RGBColor = hex_to_rgb_color(accent_color)

    for field_name in LetterData.model_fields:
        if field_name in _NON_SECTION_FIELDS:
            continue
        renderer = _SECTION_RENDERERS.get(field_name)
        if renderer is None:
            raise RuntimeError(
                f"office_export.letter_docx: LetterData field {field_name!r} has no "
                "registered renderer in _SECTION_RENDERERS — a schema field would "
                "silently go unexported from the .docx writer."
            )
        renderer(document, letter, labels, color)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
