# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""Shared ``python-docx`` foundation for the office-export writers (ADR-079, E057).

Both document kinds this epic ships — the CV writer (US296, ``cv_docx.py``) and
the cover-letter writer (US297, ``letter_docx.py``) — need the same low-level
building blocks: page setup, a base font, paragraph helpers for headings, body
text and list items, and a generic Pydantic leaf-field walker
(``_iter_leaf_paths``, moved here from ``cv_docx.py`` once a second writer
needed it — ADR-066, one implementation) for each writer's own nested-leaf
coverage guard. This module owns exactly that shared mechanics and nothing
about either document kind's *content* — no knowledge of ``TailoredCVData``,
``LetterData``, section names or field order lives here, so both writers
import it unchanged.

Deliberately **not** here, per ADR-079 clause 2/3 (US296 AC): tables, text
boxes, positioned frames, borders or shading of any kind. The accent colour
(ADR-026) is applied as a plain character-RUN colour on headings via
``add_heading`` and nowhere else — never a border, a shaded box or a
positioned frame, which is what the ADR-079 spike measured to destroy the
``creative_sidebar`` and ``executive`` presentation templates when converted.
"""

import re
import types
import typing
from typing import Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from pydantic import BaseModel

# A4 (210mm x 297mm) — the DACH-standard page size the PDF path already
# renders at (``_html_to_pdf``'s Playwright call uses ``format="A4"``); the
# docx export follows the same physical page so the two artefacts read the
# same on paper even though their layout engines differ.
PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)

# A conventional CV margin — narrower than Word's 2.54cm default so a
# realistic one/two-page CV has room to breathe without the page norms being
# fought at the margin.
PAGE_MARGIN = Mm(20)

BASE_FONT_NAME = "Calibri"
BASE_FONT_SIZE = Pt(11)

_HEX_RE = re.compile(r"^#?([0-9a-fA-F]{6})$")

_HEADING_STYLES = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}


def new_document() -> DocxDocument:
    """A fresh ``python-docx`` ``Document`` with A4 page size, standard
    margins and the base font set on the ``Normal`` style — the shared
    starting point every office-export writer builds on.
    """
    document = Document()

    section = document.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN

    normal = document.styles["Normal"]
    normal.font.name = BASE_FONT_NAME
    normal.font.size = BASE_FONT_SIZE

    return document


def hex_to_rgb_color(hex_color: str | None) -> RGBColor:
    """Parse ``"#rrggbb"`` or ``"rrggbb"`` into an :class:`RGBColor`.

    A malformed or missing value falls back to black rather than raising — a
    stale or corrupt stored accent hex must degrade the heading to plain
    black text, not crash the whole export.
    """
    match = _HEX_RE.match((hex_color or "").strip())
    if not match:
        return RGBColor(0, 0, 0)
    value = match.group(1)
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _is_blank(text: str | None) -> bool:
    return text is None or not text.strip()


def add_heading(
    document: DocxDocument, text: str | None, level: int, color: RGBColor
) -> Paragraph | None:
    """Add a ``Heading{level}`` paragraph with ``text`` painted in ``color``
    as a character-RUN colour — never a border, shading or positioned box.

    Adds nothing and returns ``None`` for blank ``text`` (``None``, ``""`` or
    whitespace-only): an empty section must produce no heading at all, never
    a heading with nothing under it.
    """
    if level not in _HEADING_STYLES:
        raise ValueError(f"add_heading: level must be 1, 2 or 3, got {level!r}")
    if _is_blank(text):
        return None

    paragraph = document.add_paragraph(style=_HEADING_STYLES[level])
    run = paragraph.add_run(text.strip())
    run.font.color.rgb = color
    return paragraph


def add_paragraph(
    document: DocxDocument,
    text: str | None,
    *,
    bold: bool = False,
    italic: bool = False,
) -> Paragraph | None:
    """Add a body paragraph (``Normal`` style). Adds nothing and returns
    ``None`` for blank ``text``."""
    if _is_blank(text):
        return None

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.italic = italic
    return paragraph


def add_bullet(document: DocxDocument, text: str | None) -> Paragraph | None:
    """Add one bulleted list item using the built-in ``List Bullet``
    paragraph style — a real Word list style, never a manually typed ``-``
    or ``•`` glyph. Adds nothing and returns ``None`` for blank
    ``text``."""
    if _is_blank(text):
        return None

    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text.strip())
    return paragraph


# ---------------------------------------------------------------------------
# Nested-field schema walker (coverage-guard infrastructure; ADR-066). Moved
# here from ``cv_docx.py`` (US297/E057 task 1.4) once a second document kind
# needed it: it was already document-kind agnostic — no reference to
# ``TailoredCVData``, ``LetterData`` or any section name — so the move is a
# pure relocation, not a rewrite. Namespace-level so every writer module AND
# its test file can import it directly against the live schema in question,
# never a hand-typed mirror of it.
# ---------------------------------------------------------------------------


def _unwrap_optional(annotation):
    """`X | None` / `Optional[X]` -> `X`. Anything else is returned unchanged."""
    origin = typing.get_origin(annotation)
    if origin in (typing.Union, types.UnionType):
        args = [a for a in typing.get_args(annotation) if a is not type(None)]
        if len(args) == 1:
            return args[0]
    return annotation


def _iter_leaf_paths(model_cls: type[BaseModel], prefix: str = "") -> Iterator[str]:
    """Walk `model_cls`'s Pydantic fields recursively, yielding one path per
    LEAF field. A field whose type is a nested `BaseModel` (directly, or as
    the element type of a `list[...]`) is descended into rather than counted
    as a leaf itself — so a field added to a nested model shows up in this
    set exactly as a top-level field of `model_cls` would, and cannot be
    missed by only checking `model_cls.model_fields`. A `list[str]`-style
    field is ONE leaf: its elements have no further schema to descend into.

    Path shape: `"parent.child"`, `"list_field[].nested_field"`,
    `"list_field[].nested_list[].leaf"`.
    """
    for name, field in model_cls.model_fields.items():
        path = f"{prefix}{name}"
        annotation = _unwrap_optional(field.annotation)
        origin = typing.get_origin(annotation)

        if origin is list:
            args = typing.get_args(annotation)
            inner = _unwrap_optional(args[0]) if args else None
            if isinstance(inner, type) and issubclass(inner, BaseModel):
                yield from _iter_leaf_paths(inner, prefix=f"{path}[].")
                continue
            yield path
            continue

        if isinstance(annotation, type) and issubclass(annotation, BaseModel):
            yield from _iter_leaf_paths(annotation, prefix=f"{path}.")
            continue

        yield path
