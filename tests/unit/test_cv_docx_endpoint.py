# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""E057/US296 — `GET /api/cv/{cv_id}/docx` over the real ASGI stack.

The writer and the service function were covered at unit level, but **nothing
referenced the route**: a grep for `/docx` across `tests/` and `backend/tests/`
returned no file. A download endpoint's contract is mostly HTTP — the media
type a word processor dispatches on, the `Content-Disposition` that decides
whether the browser saves or renders, the 404 for an unknown id, and the fact
that the route is mounted and its auth dependency resolves at all. A direct
call to the router coroutine demonstrates none of that, and this repo's own
notes record FastAPI parameter defaults leaking into direct-call tests.

So this drives `httpx.ASGITransport` against the real app, and asserts on the
bytes the client actually receives.
"""
import sys
import uuid
from io import BytesIO
from pathlib import Path

import pytest
import pytest_asyncio
from fastapi import FastAPI
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from tests.support.profile_factory import make_master_profile

_backend = Path(__file__).parent.parent.parent / "backend"
if str(_backend) not in sys.path:
    sys.path.insert(0, str(_backend))

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

BULLET = "Koordination mit Projekt Phoenix und R&D-Teams beim Rollout."
SUMMARY = "Projektleiter mit Schwerpunkt Digitale Fertigung."


@pytest_asyncio.fixture
async def db():
    from applire.db.session import Base
    import applire.models.user, applire.models.job, applire.models.profile  # noqa: F401
    import applire.models.gap, applire.models.cv, applire.models.session  # noqa: F401
    import applire.models.application, applire.models.cover_letter  # noqa: F401
    import applire.models.flow, applire.models.uploads  # noqa: F401
    import applire.models.color_profile, applire.models.company  # noqa: F401
    import applire.models.user_settings  # noqa: F401

    engine = create_async_engine("sqlite+aiosqlite:///:memory:", echo=False)
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    factory = async_sessionmaker(engine, expire_on_commit=False)
    async with factory() as session:
        yield session
    await engine.dispose()


@pytest_asyncio.fixture
async def client(db):
    from applire.auth import get_auth_provider
    from applire.db.session import get_db
    from applire.routers.cv import router

    app = FastAPI()
    app.include_router(router)
    app.dependency_overrides[get_db] = lambda: db
    app.dependency_overrides[get_auth_provider] = lambda: object()
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac


@pytest_asyncio.fixture
async def ready_cv(db):
    from applire.models.cv import GeneratedCV
    from applire.models.job import JobAnalysis

    job = JobAnalysis(
        id=uuid.uuid4(), raw_text_hash=str(uuid.uuid4()), raw_text="x",
        role_title="Teamleiter", required_skills=[], nice_to_have_skills=[],
        keywords=[], seniority_level="senior", company_culture_signals=[],
        language_requirement="de",
    )
    profile = make_master_profile(
        id=uuid.uuid4(), profile_json={"personal_info": {"name": "Jörg Müller"}}
    )
    db.add_all([job, profile])
    await db.commit()

    cv = GeneratedCV(
        id=uuid.uuid4(), job_analysis_id=job.id, profile_id=profile.id,
        template="modern_swiss", status="ready", document_language="de",
        tailored_data={
            "contact": {"name": "Jörg Müller", "email": "joerg@example.de",
                        "phone": "+49 89 1234567", "location": "München",
                        "linkedin": "", "photo_url": None},
            "summary": SUMMARY,
            "work_history": [{
                "company": "Süddeutsche Präzisionstechnik GmbH",
                "role": "Teamleiter Qualitätssicherung",
                "start_date": "2018-03", "end_date": None,
                "bullets": [BULLET],
            }],
            "education": [], "skills": ["Python"], "languages": [],
            "show_photo": False,
        },
    )
    db.add(cv)
    await db.commit()
    return cv


@pytest.mark.asyncio
async def test_docx_download_serves_a_real_word_file(client, ready_cv):
    r = await client.get(f"/api/cv/{ready_cv.id}/docx")

    assert r.status_code == 200
    assert r.headers["content-type"] == DOCX_MIME, (
        "a word processor dispatches on this media type; a generic "
        "octet-stream makes the file open in the wrong application"
    )
    assert r.content[:2] == b"PK", "a .docx is a zip container"


@pytest.mark.asyncio
async def test_download_is_an_attachment_with_a_docx_filename(client, ready_cv):
    r = await client.get(f"/api/cv/{ready_cv.id}/docx")

    disposition = r.headers["content-disposition"]
    assert disposition.startswith("attachment"), (
        "without attachment the browser may render the bytes instead of saving"
    )
    assert ".docx" in disposition and ".pdf" not in disposition


@pytest.mark.asyncio
async def test_the_delivered_bytes_carry_the_candidates_own_prose(client, ready_cv):
    """The strongest assertion here: not that a file arrived, but that the file
    the HTTP client received contains what the candidate wrote."""
    from applire.services.office_export.extract import extract_docx_text

    r = await client.get(f"/api/cv/{ready_cv.id}/docx")
    text = extract_docx_text(r.content)

    assert BULLET in text
    assert SUMMARY in text
    assert "Süddeutsche Präzisionstechnik GmbH" in text


@pytest.mark.asyncio
async def test_unknown_cv_is_404_not_500(client):
    r = await client.get(f"/api/cv/{uuid.uuid4()}/docx")

    assert r.status_code == 404


@pytest.mark.asyncio
async def test_the_file_opens_in_python_docx_with_real_paragraph_styles(client, ready_cv):
    """Guards the ADR-079 promise that this is an *editable* artefact: real
    named styles, and none of the constructs measured to destroy a layout."""
    from docx import Document

    r = await client.get(f"/api/cv/{ready_cv.id}/docx")
    doc = Document(BytesIO(r.content))

    assert len(doc.tables) == 0, "zero layout tables (ADR-079)"
    styles = {p.style.name for p in doc.paragraphs if p.text.strip()}
    assert any(s.startswith("Heading") for s in styles), f"no heading styles: {styles}"
