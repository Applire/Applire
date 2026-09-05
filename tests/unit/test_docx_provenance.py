# Copyright (C) 2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""ADR-085 clause 5 — the `.docx` export's half of the mark, and its Title.

The `.docx` is a second *delivered* artefact (ADR-079 cl. 2), so Art. 50(2)
reaches it on its own account. The carrier differs because the format does:
`docProps/custom.xml` custom document properties, plus the human-readable
duplicate in `docProps/core.xml`.

Three things are asserted that a "does the file still open" check would not see:

* the part is **discoverable** — an ``[Content_Types].xml`` override AND a
  package relationship, not merely a file in the zip. python-docx's reader walks
  the relationship graph; a part reachable by neither is invisible to every
  conforming consumer while looking perfectly present in a zip listing.
* the ECMA-376 ``fmtid``/``pid`` contract holds. LibreOffice was measured NOT to
  validate ``pid`` at all (it accepted duplicates and a zero base), so nothing
  in this environment enforces it — which is precisely why it is asserted here.
* the **Title** is language-correct and matches what the PDF template renders
  (writer collector #601: `office_export` set no core properties at all, so a
  delivered `.docx` carried an empty Title while the PDF carried a proper one).

The independent-implementation check — headless LibreOffice opening the file —
is already a hard gate in ``tests/ats/test_office_export.py`` (it converts the
export to PDF to measure page counts). That instrument covers this change for
free; it is not duplicated here (ADR-066).
"""
import io
import zipfile

import docx
import pytest
from lxml import etree

from applire.schemas.cover_letter import LetterData
from applire.schemas.cv import TailoredCVData
from applire.services.office_export.cv_docx import render_cv_docx
from applire.services.office_export.letter_docx import render_letter_docx
from applire.services.office_export.provenance import (
    CUSTOM_PROP_NAMES,
    CUSTOM_PROPS_FMTID,
    CUSTOM_PROPS_NS,
    CUSTOM_PROPS_PARTNAME,
    PROP_GENERATED,
    PROP_SOURCE_TYPE,
    read_document_provenance,
)
from applire.services.pdf_provenance import DIGITAL_SOURCE_TYPE

_CV = TailoredCVData.model_validate(
    {
        "contact": {"name": "Anna Bauer", "email": "anna@example.de"},
        "summary": "Erfahrene Ingenieurin.",
        "work_history": [
            {
                "company": "Beispiel GmbH",
                "role": "Ingenieurin",
                "start_date": "2019-01",
                "end_date": None,
                "bullets": ["Verantwortung für die Serienfertigung."],
            }
        ],
        "skills": ["Python", "Projektmanagement"],
    }
)

_LETTER = LetterData.model_validate(
    {
        "header": {"name": "Anna Bauer", "email": "anna@example.de"},
        "recipient": {"company": "Beispiel GmbH"},
        "body": {"paragraphs": ["Sehr geehrte Damen und Herren,", "Bewerbung."]},
        "signature": {"closing": "Mit freundlichen Grüßen", "name": "Anna Bauer"},
    }
)

_CASES = {
    "cv": lambda lang: render_cv_docx(_CV, lang=lang, accent_color="#123456"),
    "letter": lambda lang: render_letter_docx(_LETTER, lang=lang, accent_color="#123456"),
}

_EXPECTED_TITLES = {
    ("cv", "de"): "Lebenslauf – Anna Bauer",
    ("cv", "en"): "Curriculum Vitae – Anna Bauer",
    ("letter", "de"): "Bewerbung – Anna Bauer",
    ("letter", "en"): "Application – Anna Bauer",
}


@pytest.mark.parametrize("lang", ["de", "en"])
@pytest.mark.parametrize("kind", sorted(_CASES))
def test_every_export_carries_the_provenance_mark(kind, lang):
    properties = read_document_provenance(_CASES[kind](lang))
    assert set(properties) == set(CUSTOM_PROP_NAMES), f"{kind}/{lang}: {properties}"
    assert properties[PROP_GENERATED] == "true"
    assert properties[PROP_SOURCE_TYPE] == DIGITAL_SOURCE_TYPE


@pytest.mark.parametrize("lang", ["de", "en"])
@pytest.mark.parametrize("kind", sorted(_CASES))
def test_every_export_carries_a_language_correct_title(kind, lang):
    """Writer collector #601 — the PDF identified itself and the `.docx` did not."""
    document = docx.Document(io.BytesIO(_CASES[kind](lang)))
    assert document.core_properties.title == _EXPECTED_TITLES[(kind, lang)]
    assert "AI-generated" in (document.core_properties.comments or "")
    assert document.core_properties.keywords == "AI-generated"


@pytest.mark.parametrize("kind", sorted(_CASES))
def test_the_custom_properties_part_is_discoverable_not_merely_present(kind):
    blob = _CASES[kind]("de")
    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        names = archive.namelist()
        assert CUSTOM_PROPS_PARTNAME.lstrip("/") in names
        content_types = archive.read("[Content_Types].xml").decode()
        package_rels = archive.read("_rels/.rels").decode()

    assert 'PartName="/docProps/custom.xml"' in content_types
    assert "custom-properties+xml" in content_types
    assert "custom-properties" in package_rels
    assert 'Target="docProps/custom.xml"' in package_rels or (
        'Target="/docProps/custom.xml"' in package_rels
    )


@pytest.mark.parametrize("kind", sorted(_CASES))
def test_custom_properties_honour_the_ecma376_fmtid_and_pid_contract(kind):
    with zipfile.ZipFile(io.BytesIO(_CASES[kind]("de"))) as archive:
        root = etree.fromstring(
            archive.read(CUSTOM_PROPS_PARTNAME.lstrip("/")),
            parser=etree.XMLParser(resolve_entities=False, no_network=True),
        )

    entries = root.findall(f"{{{CUSTOM_PROPS_NS}}}property")
    assert [e.get("fmtid") for e in entries] == [CUSTOM_PROPS_FMTID] * len(entries)
    # pid is 1-based with 0 and 1 reserved, and must be unique and ascending.
    assert [int(e.get("pid")) for e in entries] == list(
        range(2, 2 + len(entries))
    ), "ECMA-376 pid contract violated — no consumer here enforces it, so we do"


@pytest.mark.parametrize("kind", sorted(_CASES))
def test_the_export_still_opens_and_keeps_its_content(kind):
    document = docx.Document(io.BytesIO(_CASES[kind]("de")))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Anna Bauer" in text


@pytest.mark.parametrize("lang", ["de", "en"])
@pytest.mark.parametrize("kind", sorted(_CASES))
def test_export_carries_no_provider_property(kind, lang):
    """ADR-085 clause 3, amended by founder ruling 15 (2026-09-05).

    Earlier the `.docx` mark named the configured LLM provider family as a
    custom property. Ruling 15 dropped that property on both carriers.

    The expected set below is spelled out **independently** of
    ``CUSTOM_PROP_NAMES`` rather than compared against it: a mutation that
    reintroduces a provider property consistently in both the constant and
    the writer would sail through an ``== set(CUSTOM_PROP_NAMES)`` check
    (both sides move together), so the ground truth here is hardcoded.
    """
    expected_properties = {
        "AIGenerated",
        "AIGeneratedBy",
        "AIGeneratedAt",
        "AIDigitalSourceType",
        "AIMarkingSpec",
    }
    properties = read_document_provenance(_CASES[kind](lang))
    assert set(properties) == expected_properties, properties
    assert set(properties) == set(CUSTOM_PROP_NAMES), properties

    blob = _CASES[kind](lang)
    for forbidden in (b"mistral", b"openrouter", b"openai"):
        assert forbidden not in blob, forbidden


def test_read_document_provenance_says_nothing_for_an_unmarked_docx():
    """Negative control — the detector must be able to say "no"."""
    from docx import Document

    buffer = io.BytesIO()
    Document().save(buffer)
    assert read_document_provenance(buffer.getvalue()) == {}


def test_the_docx_extractor_the_ats_audit_uses_is_unaffected():
    """The ADR-039 audit reads the export back through this function."""
    from applire.services.office_export.extract import extract_docx_text

    assert "Anna Bauer" in extract_docx_text(_CASES["cv"]("de"))
