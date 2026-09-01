# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""#640 — the CV import path drops every Word table, so a *tabellarischer
Lebenslauf* (the standard DACH layout) reaches the vault with its entire
employment and education history missing.

``cv_parser.extract_text_from_docx`` read only ``document.paragraphs``. Word
table cells are not paragraphs of the document body, so every fact laid out in
a table was discarded **before** the extraction LLM was called — no prompt
change could recover it. The section *headings* survived (ordinary paragraphs),
which made the result look structured while carrying nothing underneath.

The fix delegates to ``office_export.extract.extract_docx_text``, the
document-order walker built for E057/ADR-079, and keeps the import seam's own
blank-line normalisation so the whitespace shape the extraction prompt sees is
unchanged. That also settles an ADR-066 tension: two implementations of one
logical operation — "get the text out of a ``.docx``" — one measurably
deficient.

The measured behavioural delta on documents WITHOUT tables is exactly one
character class: a literal tab becomes a single space
(``test_tab_stop_becomes_a_space``). Everything else is byte-identical, which
is pinned here rather than left to trust — the import path feeds the extraction
prompts, so a silent change in whitespace shape is a change in model input.
"""
from io import BytesIO

import pytest

docx = pytest.importorskip("docx", reason="python-docx is required for DOCX import")
from docx import Document  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import nsdecls  # noqa: E402

from applire.services.cv_parser import extract_text_from_docx  # noqa: E402


def _bytes(build) -> bytes:
    d = Document()
    build(d)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def _tabellarischer_lebenslauf(d) -> None:
    """The #640 reproduction: headings as paragraphs, facts in 2-column tables."""
    d.add_paragraph("Jörg Müller")
    d.add_paragraph("joerg.mueller@example.de · +49 89 1234567")
    d.add_paragraph("Berufserfahrung")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "03/2018 – heute"
    t.cell(0, 1).text = "Teamleiter Qualitätssicherung, Süddeutsche Präzisionstechnik GmbH"
    t.cell(1, 0).text = "09/2014 – 02/2018"
    t.cell(1, 1).text = "Ingenieur Fertigungsplanung, Nordwerk Systeme AG"
    d.add_paragraph("Ausbildung")
    t2 = d.add_table(rows=1, cols=2)
    t2.cell(0, 0).text = "10/2006 – 03/2011"
    t2.cell(0, 1).text = "Dipl.-Ing. Maschinenbau, Technische Universität München"
    d.add_paragraph("Kenntnisse")
    d.add_paragraph("Python, Six Sigma, SPC")


# Every fact the reproduction puts in a table. Nine of nine were lost.
TABLE_BORNE_FACTS = [
    "Teamleiter Qualitätssicherung",
    "Süddeutsche Präzisionstechnik GmbH",
    "Ingenieur Fertigungsplanung",
    "Nordwerk Systeme AG",
    "03/2018 – heute",
    "09/2014 – 02/2018",
    "10/2006 – 03/2011",
    "Dipl.-Ing. Maschinenbau",
    "Technische Universität München",
]


def test_fixture_really_uses_tables():
    """Guard the guard: the reproduction must fail for the reason claimed.

    If a python-docx change ever made ``add_table`` emit paragraphs, every
    assertion below would pass while testing nothing.
    """
    d = Document()
    _tabellarischer_lebenslauf(d)

    assert len(d.tables) == 2, "the fixture must lay its facts out in real tables"
    paragraph_text = "\n".join(p.text for p in d.paragraphs)
    for fact in TABLE_BORNE_FACTS:
        assert fact not in paragraph_text, (
            f"{fact!r} must live in a table cell, not a body paragraph — "
            "otherwise this file does not reproduce #640"
        )


@pytest.mark.parametrize("fact", TABLE_BORNE_FACTS)
def test_table_borne_fact_reaches_the_extracted_text(fact):
    """#640: nine of nine of these were silently dropped before the LLM ran."""
    text = extract_text_from_docx(_bytes(_tabellarischer_lebenslauf))

    assert fact in text, (
        f"{fact!r} is laid out in a Word table and never reached the vault"
    )


def test_table_content_keeps_document_order():
    """Not merely present — present *where the reader sees it*.

    ``document.paragraphs`` + ``document.tables`` concatenated would report both
    tables AFTER 'Kenntnisse'; the extraction prompt would then read the
    Ausbildung dates as belonging to the skills section.
    """
    text = extract_text_from_docx(_bytes(_tabellarischer_lebenslauf))

    assert text.index("Berufserfahrung") < text.index("Teamleiter Qualitätssicherung")
    assert text.index("Teamleiter Qualitätssicherung") < text.index("Ausbildung")
    assert text.index("Ausbildung") < text.index("Dipl.-Ing. Maschinenbau")
    assert text.index("Dipl.-Ing. Maschinenbau") < text.index("Kenntnisse")


# A minimal DrawingML text box, built the way ``test_office_export_extract.py``
# builds one — python-docx exposes no public API for them, and hand-assembling
# the elements via ``qn()`` silently produces a malformed shape.
_TXBX_EXTRA_NS = (
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
    'xmlns:v="urn:schemas-microsoft-com:vml"'
)


def _text_box_xml(text: str) -> str:
    return (
        '<w:drawing %s %s>'
        '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        '<wp:extent cx="2000000" cy="500000"/>'
        '<wp:docPr id="1" name="TextBox"/>'
        '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp><wps:cNvSpPr txBox="1"/>'
        '<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="2000000" cy="500000"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></wps:spPr>'
        f'<wps:txbx><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:txbxContent></wps:txbx>'
        '<wps:bodyPr/></wps:wsp>'
        '</a:graphicData></a:graphic></wp:inline></w:drawing>'
    ) % (nsdecls("w", "wp", "a"), _TXBX_EXTRA_NS)


def test_text_box_content_reaches_the_extracted_text():
    """A text box is neither a body paragraph nor a table, so both of the old
    reader's accessors missed it. Word CV templates use them for contact blocks
    and side columns."""

    def build(d):
        d.add_paragraph("Berufserfahrung")
        p = d.add_paragraph()
        p.add_run()._r.append(parse_xml(_text_box_xml("joerg.mueller@example.de")))
        d.add_paragraph("Kenntnisse")

    text = extract_text_from_docx(_bytes(build))

    assert "joerg.mueller@example.de" in text
    assert text.index("Berufserfahrung") < text.index("joerg.mueller@example.de")
    assert text.index("joerg.mueller@example.de") < text.index("Kenntnisse")


# ---------------------------------------------------------------------------
# The other half of the contract: what must NOT change.
#
# This path feeds the extraction prompts, so any drift in whitespace shape is
# drift in model input. These pin the measured no-change result.
# ---------------------------------------------------------------------------


def test_paragraph_only_document_is_unchanged():
    """Byte-for-byte identical to what the paragraphs reader produced."""

    def build(d):
        for p in [
            "Jörg Müller",
            "joerg.mueller@example.de · +49 89 1234567",
            "",
            "Berufserfahrung",
            "Teamleiter Qualitätssicherung, Süddeutsche Präzisionstechnik GmbH",
            "03/2018 – heute",
            "",
            "Kenntnisse",
            "Python, Six Sigma, SPC",
        ]:
            d.add_paragraph(p)

    assert extract_text_from_docx(_bytes(build)) == (
        "Jörg Müller\n"
        "joerg.mueller@example.de · +49 89 1234567\n"
        "Berufserfahrung\n"
        "Teamleiter Qualitätssicherung, Süddeutsche Präzisionstechnik GmbH\n"
        "03/2018 – heute\n"
        "Kenntnisse\n"
        "Python, Six Sigma, SPC"
    )


def test_blank_paragraphs_stay_collapsed():
    """The walker newline-terminates EVERY paragraph including empty ones; the
    import seam's normalisation must keep dropping them, or a spacing-heavy CV
    reaches the prompt padded with blank lines."""

    def build(d):
        d.add_paragraph("Berufserfahrung")
        for _ in range(4):
            d.add_paragraph("")
        d.add_paragraph("Teamleiter Qualitätssicherung")

    assert extract_text_from_docx(_bytes(build)) == (
        "Berufserfahrung\nTeamleiter Qualitätssicherung"
    )


def test_whitespace_only_paragraphs_stay_collapsed():
    """A paragraph of spaces is as empty as an empty one, and the old reader's
    ``if p.text.strip()`` dropped it."""

    def build(d):
        d.add_paragraph("Berufserfahrung")
        d.add_paragraph("   ")
        d.add_paragraph("Teamleiter Qualitätssicherung")

    assert extract_text_from_docx(_bytes(build)) == (
        "Berufserfahrung\nTeamleiter Qualitätssicherung"
    )


def test_tab_stop_becomes_a_space():
    """The one measured behavioural change on a table-free document, pinned so
    it is a decision rather than an accident.

    Word CV templates right-align a date with a tab stop. The walker emits one
    space for ``<w:tab/>``; the paragraphs reader emitted a literal tab. Neither
    loses a fact, and a space is the better token boundary for the extraction
    prompt.
    """

    def build(d):
        p = d.add_paragraph("Teamleiter Qualitätssicherung")
        p.add_run("\t")
        p.add_run("03/2018 – heute")

    text = extract_text_from_docx(_bytes(build))

    assert text == "Teamleiter Qualitätssicherung 03/2018 – heute"
    assert "\t" not in text


def test_empty_document_still_returns_empty_string():
    assert extract_text_from_docx(_bytes(lambda d: None)) == ""
