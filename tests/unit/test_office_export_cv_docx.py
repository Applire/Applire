# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""E057 / US296 (ADR-079) — the direct python-docx CV writer.

Two layers under test:
  * ``office_export._common`` — the shared docx foundation (page setup, base
    font, heading/paragraph/bullet helpers, accent-colour parsing). CV-kind
    agnostic; the cover-letter writer (US297) will import it unchanged.
  * ``office_export.cv_docx`` — the pure ``TailoredCVData -> bytes`` writer.

No HTML, no template engine, no subprocess anywhere in this file's import
graph (ADR-079 clause 2) — a test importing Jinja or shelling out to
LibreOffice would itself be a boundary violation.
"""

import io
import zipfile

from docx.document import Document as OpenDocument
from docx.shared import RGBColor

from applire.schemas.cv import (
    TailoredCVData,
    TailoredCertification,
    TailoredContact,
    TailoredEducationEntry,
    TailoredLanguage,
    TailoredProjectEntry,
    TailoredWorkEntry,
)
from applire.templates.labels import cv_labels


# ---------------------------------------------------------------------------
# _common.py — shared docx foundation
# ---------------------------------------------------------------------------


class TestNewDocument:
    def test_returns_a_docx_document(self):
        from applire.services.office_export._common import new_document

        document = new_document()
        assert isinstance(document, OpenDocument)

    def test_sets_a4_page_size(self):
        from applire.services.office_export._common import new_document

        document = new_document()
        section = document.sections[0]
        # OOXML stores page size in twips; an Mm->twips->Emu round trip
        # introduces a few hundred EMU of rounding (~0.01mm) — compare in mm,
        # rounded, rather than asserting exact Emu equality.
        assert round(section.page_width.mm) == 210
        assert round(section.page_height.mm) == 297

    def test_sets_base_font_on_normal_style(self):
        from applire.services.office_export._common import (
            BASE_FONT_NAME,
            BASE_FONT_SIZE,
            new_document,
        )

        document = new_document()
        normal = document.styles["Normal"]
        assert normal.font.name == BASE_FONT_NAME
        assert normal.font.size == BASE_FONT_SIZE


class TestHexToRgbColor:
    def test_parses_hash_prefixed_hex(self):
        from applire.services.office_export._common import hex_to_rgb_color

        assert hex_to_rgb_color("#1a73e8") == RGBColor(0x1A, 0x73, 0xE8)

    def test_parses_hex_without_hash(self):
        from applire.services.office_export._common import hex_to_rgb_color

        assert hex_to_rgb_color("1A73E8") == RGBColor(0x1A, 0x73, 0xE8)

    def test_malformed_value_falls_back_to_black(self):
        from applire.services.office_export._common import hex_to_rgb_color

        assert hex_to_rgb_color("not-a-color") == RGBColor(0, 0, 0)

    def test_none_falls_back_to_black(self):
        from applire.services.office_export._common import hex_to_rgb_color

        assert hex_to_rgb_color(None) == RGBColor(0, 0, 0)


class TestAddHeading:
    def test_creates_named_heading_style_paragraph(self):
        from applire.services.office_export._common import add_heading, new_document

        document = new_document()
        add_heading(document, "Berufserfahrung", 2, RGBColor(0, 0, 0))

        assert document.paragraphs[-1].style.name == "Heading 2"
        assert document.paragraphs[-1].text == "Berufserfahrung"

    def test_applies_accent_color_as_run_color(self):
        from applire.services.office_export._common import add_heading, new_document

        document = new_document()
        color = RGBColor(0x1A, 0x73, 0xE8)
        add_heading(document, "Kenntnisse", 2, color)

        run = document.paragraphs[-1].runs[0]
        assert run.font.color.rgb == color

    def test_blank_text_adds_no_paragraph(self):
        from applire.services.office_export._common import add_heading, new_document

        document = new_document()
        result = add_heading(document, "", 2, RGBColor(0, 0, 0))

        assert result is None
        assert len(document.paragraphs) == 0

    def test_none_text_adds_no_paragraph(self):
        from applire.services.office_export._common import add_heading, new_document

        document = new_document()
        result = add_heading(document, None, 2, RGBColor(0, 0, 0))

        assert result is None
        assert len(document.paragraphs) == 0

    def test_whitespace_only_text_adds_no_paragraph(self):
        from applire.services.office_export._common import add_heading, new_document

        document = new_document()
        result = add_heading(document, "   ", 1, RGBColor(0, 0, 0))

        assert result is None
        assert len(document.paragraphs) == 0

    def test_invalid_level_raises(self):
        import pytest
        from applire.services.office_export._common import add_heading, new_document

        document = new_document()
        with pytest.raises(ValueError):
            add_heading(document, "Text", 4, RGBColor(0, 0, 0))


class TestAddParagraph:
    def test_adds_body_paragraph_with_normal_style(self):
        from applire.services.office_export._common import add_paragraph, new_document

        document = new_document()
        add_paragraph(document, "Berlin, Germany")

        assert document.paragraphs[-1].style.name == "Normal"
        assert document.paragraphs[-1].text == "Berlin, Germany"

    def test_bold_flag_applies_to_run(self):
        from applire.services.office_export._common import add_paragraph, new_document

        document = new_document()
        add_paragraph(document, "TechCorp GmbH", bold=True)

        assert document.paragraphs[-1].runs[0].bold is True

    def test_blank_text_adds_nothing(self):
        from applire.services.office_export._common import add_paragraph, new_document

        document = new_document()
        result = add_paragraph(document, "")

        assert result is None
        assert len(document.paragraphs) == 0

    def test_none_text_adds_nothing(self):
        from applire.services.office_export._common import add_paragraph, new_document

        document = new_document()
        result = add_paragraph(document, None)

        assert result is None
        assert len(document.paragraphs) == 0


class TestAddBullet:
    def test_uses_list_bullet_style(self):
        from applire.services.office_export._common import add_bullet, new_document

        document = new_document()
        add_bullet(document, "Python")

        assert document.paragraphs[-1].style.name == "List Bullet"
        assert document.paragraphs[-1].text == "Python"

    def test_blank_text_adds_nothing(self):
        from applire.services.office_export._common import add_bullet, new_document

        document = new_document()
        result = add_bullet(document, "  ")

        assert result is None
        assert len(document.paragraphs) == 0


# ---------------------------------------------------------------------------
# cv_docx.py — the pure TailoredCVData -> bytes writer
# ---------------------------------------------------------------------------

ACCENT = "#1a73e8"


def _open(docx_bytes: bytes):
    from docx import Document as open_docx

    return open_docx(io.BytesIO(docx_bytes))


def _all_text(docx_bytes: bytes) -> str:
    document = _open(docx_bytes)
    return "\n".join(p.text for p in document.paragraphs)


def _test_photo_bytes() -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (10, 10), color=(200, 30, 30)).save(buf, format="PNG")
    return buf.getvalue()


def _full_tailored_cv() -> TailoredCVData:
    """One populated instance of every TailoredCVData field, with distinct
    MARKER_* strings per field so content tests can assert precisely on
    which section produced which text."""
    return TailoredCVData(
        contact=TailoredContact(
            name="Erika Mustermann",
            email="erika@example.com",
            phone="+49 30 1234567",
            location="Berlin, Germany",
            linkedin="linkedin.com/in/erikamustermann",
        ),
        summary="MARKER_SUMMARY_TEXT experienced backend engineer.",
        work_history=[
            TailoredWorkEntry(
                company="MARKER_COMPANY_ALPHA",
                role="MARKER_ROLE_ALPHA",
                start_date="2020-01",
                end_date="2022-06",
                bullets=["MARKER_BULLET_ALPHA_1", "MARKER_BULLET_ALPHA_2"],
                team_size=42,
                budget_managed="MARKER_BUDGET_ALPHA",
                industry_context="MARKER_INDUSTRY_ALPHA",
                projects=[
                    TailoredProjectEntry(
                        name="MARKER_NESTED_PROJECT_ALPHA",
                        bullets=["MARKER_NESTED_PROJECT_BULLET_ALPHA"],
                    )
                ],
            ),
            TailoredWorkEntry(
                company="MARKER_COMPANY_BETA",
                role="MARKER_ROLE_BETA",
                start_date="2022-07",
                end_date=None,
                bullets=["MARKER_BULLET_BETA_1"],
            ),
        ],
        skills=["MARKER_SKILL_PYTHON", "MARKER_SKILL_SQL"],
        education=[
            TailoredEducationEntry(
                institution="MARKER_UNIVERSITY",
                degree="MARKER_DEGREE",
                field="MARKER_FIELD",
                start_date="2016",
                end_date="2019",
            )
        ],
        languages=[
            TailoredLanguage(language="MARKER_LANGUAGE_GERMAN", level="MARKER_LEVEL_NATIVE")
        ],
        projects=[
            TailoredProjectEntry(
                name="MARKER_STANDALONE_PROJECT",
                bullets=["MARKER_STANDALONE_PROJECT_BULLET"],
            )
        ],
        certifications=[
            TailoredCertification(
                name="MARKER_CERT_NAME",
                issuing_organization="MARKER_CERT_ORG",
                date_obtained="2021-01",
                expiry_date="2024-01",
            )
        ],
        show_photo=True,
    )


class TestRenderCvDocxStructure:
    def test_returns_bytes(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_bytes_are_a_valid_docx(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        document = _open(result)
        assert len(document.paragraphs) > 0

    def test_no_tables(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        document = _open(result)
        assert document.tables == []

    def test_no_layout_tables_text_boxes_or_positioned_frames_in_xml(self):
        """Direct XML inspection (US296 AC): zero tables, zero text boxes,
        zero positioned frames — the exact constructs the ADR-079 spike
        measured to destroy creative_sidebar/executive. Namespace-agnostic
        local-name check, not a raw substring search, so a namespace
        *declaration* (harmless boilerplate on the root element) cannot
        produce a false positive."""
        from lxml import etree

        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        document = _open(result)
        local_names = {etree.QName(el).localname for el in document.element.body.iter()}

        forbidden = {"tbl", "pBdr", "framePr", "pict", "txbx", "shape", "textbox"}
        found = local_names & forbidden
        assert not found, f"forbidden OOXML constructs present: {found}"

    def test_uses_heading_1_2_and_3_styles(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        document = _open(result)
        style_names = {p.style.name for p in document.paragraphs}
        assert {"Heading 1", "Heading 2", "Heading 3"} <= style_names

    def test_uses_list_bullet_style(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        document = _open(result)
        style_names = {p.style.name for p in document.paragraphs}
        assert "List Bullet" in style_names

    def test_accent_color_appears_only_on_heading_runs(self):
        from applire.services.office_export._common import hex_to_rgb_color
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        document = _open(result)
        color = hex_to_rgb_color(ACCENT)
        heading_styles = {"Heading 1", "Heading 2", "Heading 3"}

        saw_accent_on_heading = False
        for paragraph in document.paragraphs:
            is_heading = paragraph.style.name in heading_styles
            for run in paragraph.runs:
                if run.font.color is not None and run.font.color.rgb == color:
                    assert is_heading, (
                        f"accent colour on a non-heading run: "
                        f"style={paragraph.style.name!r} text={run.text!r}"
                    )
                    saw_accent_on_heading = True
        assert saw_accent_on_heading, "expected at least one heading run painted in accent colour"


class TestSectionCoverageGuard:
    def test_every_tailored_cv_field_has_a_registered_renderer(self):
        """The #619-style gate: the renderer dispatch is derived from
        TailoredCVData.model_fields, never a hand-written list, so a new
        schema field cannot silently go unexported."""
        from applire.services.office_export.cv_docx import (
            _NON_SECTION_FIELDS,
            _SECTION_RENDERERS,
        )

        covered = set(_SECTION_RENDERERS) | set(_NON_SECTION_FIELDS)
        assert covered == set(TailoredCVData.model_fields)


class TestNestedLeafCoverageGuard:
    """SF-PROFILE.8's lesson, applied here: TestSectionCoverageGuard above
    only sees TailoredCVData's TOP-LEVEL fields, so a field nested inside
    TailoredWorkEntry (or any other nested model) could go unrendered
    forever without turning anything red — which is exactly what happened
    to work_history[].team_size/budget_managed/industry_context (rendered
    by all seven PDF templates, silently absent from the .docx). These
    tests walk the FULL field tree, so a new nested field must be given an
    explicit rendered/not-rendered decision or the suite goes red."""

    def test_iter_leaf_paths_finds_nested_work_history_fields(self):
        """Sanity check on the walker itself, independent of the
        registries: it must actually descend into TailoredWorkEntry and its
        nested TailoredProjectEntry, not stop at 'work_history' as one leaf."""
        from applire.services.office_export.cv_docx import _iter_leaf_paths

        leaves = set(_iter_leaf_paths(TailoredCVData))
        for expected in (
            "work_history[].team_size",
            "work_history[].budget_managed",
            "work_history[].industry_context",
            "work_history[].projects[].name",
            "work_history[].projects[].bullets",
            "contact.photo_url",
        ):
            assert expected in leaves, f"walker did not find {expected!r}"
        assert "summary" in leaves  # a plain top-level field is still one leaf
        assert "skills" in leaves  # list[str] has no sub-schema to descend into

    def test_every_tailored_cv_leaf_field_is_accounted_for(self):
        """Every leaf in the TailoredCVData tree — including fields nested
        inside TailoredWorkEntry/TailoredProjectEntry/etc — must be in
        EXACTLY ONE of _RENDERED_LEAVES / _NOT_RENDERED_LEAVES. A leaf in
        neither is a silent omission (this is what let the role facts slip
        through); a leaf in both is a self-contradiction."""
        from applire.services.office_export.cv_docx import (
            _NOT_RENDERED_LEAVES,
            _RENDERED_LEAVES,
            _iter_leaf_paths,
        )

        all_leaves = set(_iter_leaf_paths(TailoredCVData))
        accounted = set(_RENDERED_LEAVES) | set(_NOT_RENDERED_LEAVES)

        assert all_leaves == accounted, (
            f"unaccounted leaves: {all_leaves - accounted!r}; "
            f"stale registry entries no longer in the schema: {accounted - all_leaves!r}"
        )

    def test_rendered_and_not_rendered_leaves_are_disjoint(self):
        from applire.services.office_export.cv_docx import (
            _NOT_RENDERED_LEAVES,
            _RENDERED_LEAVES,
        )

        overlap = set(_RENDERED_LEAVES) & set(_NOT_RENDERED_LEAVES)
        assert not overlap, f"leaves classified as both rendered and not: {overlap!r}"

    def test_not_rendered_leaves_have_written_reasons(self):
        from applire.services.office_export.cv_docx import _NOT_RENDERED_LEAVES

        assert _NOT_RENDERED_LEAVES, "expected at least one deliberately-not-rendered leaf"
        for leaf, reason in _NOT_RENDERED_LEAVES.items():
            assert isinstance(reason, str) and len(reason.strip()) >= 15, (
                f"{leaf!r} has no real written reason: {reason!r}"
            )


class TestRenderCvDocxContent:
    def test_contact_fields_appear(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        cv = _full_tailored_cv()
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        assert cv.contact.name in text
        assert cv.contact.email in text
        assert cv.contact.phone in text
        assert cv.contact.location in text
        assert cv.contact.linkedin in text

    def test_summary_appears(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert "MARKER_SUMMARY_TEXT" in text

    def test_work_history_and_nested_projects_appear(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        for marker in (
            "MARKER_COMPANY_ALPHA",
            "MARKER_ROLE_ALPHA",
            "MARKER_BULLET_ALPHA_1",
            "MARKER_BULLET_ALPHA_2",
            "MARKER_NESTED_PROJECT_ALPHA",
            "MARKER_NESTED_PROJECT_BULLET_ALPHA",
            "MARKER_COMPANY_BETA",
            "MARKER_ROLE_BETA",
            "MARKER_BULLET_BETA_1",
        ):
            assert marker in text, f"missing {marker!r}"

    def test_ongoing_role_uses_present_label(self):
        """work_history[1].end_date=None must render as labels['present']
        ('heute' in German), never as a blank or literal 'None'."""
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert cv_labels("de")["present"] in text
        assert "None" not in text

    def test_skills_appear(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert "MARKER_SKILL_PYTHON" in text
        assert "MARKER_SKILL_SQL" in text

    def test_education_appears(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert "MARKER_UNIVERSITY" in text
        assert "MARKER_DEGREE" in text
        assert "MARKER_FIELD" in text

    def test_languages_appear(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert "MARKER_LANGUAGE_GERMAN" in text
        assert "MARKER_LEVEL_NATIVE" in text

    def test_standalone_projects_appear(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert "MARKER_STANDALONE_PROJECT" in text
        assert "MARKER_STANDALONE_PROJECT_BULLET" in text

    def test_certifications_appear(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert "MARKER_CERT_NAME" in text
        assert "MARKER_CERT_ORG" in text
        assert "2021-01" in text
        assert "2024-01" in text

    def test_photo_embedded_when_photo_bytes_present(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(
            _full_tailored_cv(), lang="de", accent_color=ACCENT, photo_bytes=_test_photo_bytes()
        )
        document = _open(result)
        assert len(document.inline_shapes) >= 1

    def test_photo_absent_when_photo_bytes_none(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(
            _full_tailored_cv(), lang="de", accent_color=ACCENT, photo_bytes=None
        )
        document = _open(result)
        assert len(document.inline_shapes) == 0

    def test_photo_bytes_defaults_to_none(self):
        """photo_bytes is optional — callers with show_photo=False never pass it."""
        from applire.services.office_export.cv_docx import render_cv_docx

        result = render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT)
        document = _open(result)
        assert len(document.inline_shapes) == 0


class TestRenderCvDocxRoleFacts:
    """#328-style per-role quantified facts (team_size/budget_managed/
    industry_context) — rendered by all seven PDF templates via
    cv_labels()'s role_team_size/role_budget/role_industry slots. ADR-079's
    premise is that the export carries the same content set as the PDF, so
    these are not optional."""

    def test_role_facts_appear_when_present(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        labels_de = cv_labels("de")
        assert f"{labels_de['role_team_size']}: 42" in text
        assert f"{labels_de['role_budget']}: MARKER_BUDGET_ALPHA" in text
        assert f"{labels_de['role_industry']}: MARKER_INDUSTRY_ALPHA" in text

    def test_team_size_zero_renders(self):
        """The trap: team_size: int | None, and 0 is a valid, meaningful
        team size ("None means 'not stated' — 0 is a valid team_size",
        schemas/cv.py:140). A `if entry.team_size:` truthiness guard would
        silently drop a real zero; this pins that it does not."""
        from applire.services.office_export.cv_docx import render_cv_docx

        cv = TailoredCVData(
            contact=TailoredContact(name="Solo Kandidat"),
            work_history=[
                TailoredWorkEntry(company="OneCo", role="Founder", team_size=0),
            ],
        )
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        labels_de = cv_labels("de")
        assert f"{labels_de['role_team_size']}: 0" in text, (
            f"team_size=0 did not render — a truthiness guard would produce "
            f"exactly this miss:\n{text}"
        )

    def test_role_facts_absent_when_all_none(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        cv = TailoredCVData(
            contact=TailoredContact(name="Solo Kandidat"),
            work_history=[
                TailoredWorkEntry(
                    company="OneCo", role="Engineer",
                    team_size=None, budget_managed=None, industry_context=None,
                ),
            ],
        )
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        labels_de = cv_labels("de")
        assert labels_de["role_team_size"] not in text
        assert labels_de["role_budget"] not in text
        assert labels_de["role_industry"] not in text

    def test_role_facts_only_entry_still_counts_as_content(self):
        """A work entry whose ONLY content is team_size=0 must still count
        as 'has content': the Experience heading is not suppressed, and
        rendering does not crash."""
        from applire.services.office_export.cv_docx import render_cv_docx

        cv = TailoredCVData(
            contact=TailoredContact(name="Solo Kandidat"),
            work_history=[TailoredWorkEntry(team_size=0)],
        )
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        labels_de = cv_labels("de")
        assert labels_de["experience"] in text
        assert f"{labels_de['role_team_size']}: 0" in text

    def test_role_facts_use_en_labels_in_english(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="en", accent_color=ACCENT))
        labels_en = cv_labels("en")
        assert f"{labels_en['role_team_size']}: 42" in text
        assert cv_labels("de")["role_team_size"] not in text


class TestRenderCvDocxDegenerateInputs:
    def test_minimal_cv_does_not_crash(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        cv = TailoredCVData(contact=TailoredContact())
        result = render_cv_docx(cv, lang="de", accent_color=ACCENT)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_empty_sections_produce_no_heading(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        labels_de = cv_labels("de")
        cv = TailoredCVData(
            contact=TailoredContact(name="Solo Kandidat"),
            summary="",
            work_history=[
                TailoredWorkEntry(company="OnlyCo", role="OnlyRole", bullets=["Did a thing"]),
            ],
            skills=[],
            education=[],
            languages=[],
            projects=[],
            certifications=[],
        )
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        assert labels_de["experience"] in text  # the one populated section
        for key in ("summary", "skills", "education", "languages", "projects", "certifications"):
            assert labels_de[key] not in text, f"empty section {key!r} emitted its heading"

    def test_none_optional_fields_do_not_crash(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        cv = TailoredCVData(
            contact=TailoredContact(name="Anna Bauer", email=None, phone=None, linkedin=None),
            work_history=[TailoredWorkEntry(company="Foo", role="Bar", end_date=None)],
        )
        result = render_cv_docx(cv, lang="de", accent_color=ACCENT)
        assert isinstance(result, bytes)

    def test_blank_role_and_company_does_not_emit_empty_heading3(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        cv = TailoredCVData(
            contact=TailoredContact(name="Anna Bauer"),
            work_history=[
                TailoredWorkEntry(company="", role="", bullets=["A real bullet survives"]),
            ],
        )
        result = render_cv_docx(cv, lang="de", accent_color=ACCENT)
        document = _open(result)
        heading3_texts = [p.text for p in document.paragraphs if p.style.name == "Heading 3"]
        assert all(t.strip() for t in heading3_texts), f"found an empty Heading 3: {heading3_texts!r}"
        assert "A real bullet survives" in _all_text(result)

    def test_fully_blank_work_entry_suppresses_experience_heading(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        labels_de = cv_labels("de")
        cv = TailoredCVData(
            contact=TailoredContact(name="Anna Bauer"),
            work_history=[TailoredWorkEntry()],  # every field blank/default
        )
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        assert labels_de["experience"] not in text

    def test_blank_skill_strings_suppress_skills_heading(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        labels_de = cv_labels("de")
        cv = TailoredCVData(contact=TailoredContact(name="Anna Bauer"), skills=["", "   "])
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        assert labels_de["skills"] not in text

    def test_twenty_work_entries_do_not_crash(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        entries = [
            TailoredWorkEntry(company=f"Company{i}", role=f"Role{i}", bullets=[f"Bullet{i}"])
            for i in range(20)
        ]
        cv = TailoredCVData(contact=TailoredContact(name="Busy Bee"), work_history=entries)
        text = _all_text(render_cv_docx(cv, lang="de", accent_color=ACCENT))
        for i in range(20):
            assert f"Company{i}" in text


class TestRenderCvDocxI18n:
    def test_de_labels_used_for_section_headings(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="de", accent_color=ACCENT))
        assert cv_labels("de")["experience"] in text
        assert cv_labels("en")["experience"] not in text

    def test_en_labels_used_for_section_headings(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="en", accent_color=ACCENT))
        assert cv_labels("en")["experience"] in text
        assert cv_labels("de")["experience"] not in text

    def test_unknown_lang_falls_back_to_german(self):
        from applire.services.office_export.cv_docx import render_cv_docx

        text = _all_text(render_cv_docx(_full_tailored_cv(), lang="fr", accent_color=ACCENT))
        assert cv_labels("de")["experience"] in text
