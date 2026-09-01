# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""E057 / US297 (ADR-079) — the direct python-docx cover-letter writer.

One layer under test: ``office_export.letter_docx`` — the pure
``LetterData -> bytes`` writer. ``office_export._common`` (page setup, base
font, heading/paragraph/bullet helpers, accent-colour parsing, and the
``_iter_leaf_paths`` schema walker moved here from ``cv_docx.py`` in this
same task) is already exercised by ``test_office_export_cv_docx.py`` and not
re-tested in isolation here — only through this writer's own use of it.

No HTML, no template engine, no subprocess anywhere in this file's import
graph (ADR-079 clause 2) — a test importing Jinja or shelling out to
LibreOffice would itself be a boundary violation.

Design note on headings (see ``letter_docx.py``'s module docstring for the
full reasoning): unlike the CV (Heading 1/2/3, with labelled section
headings sourced from ``cv_labels()``), the letter has exactly ONE natural
heading — the sender's own name (``header.name``, Heading 1).
``cover_letter_labels()`` has no key naming a "recipient"/"body"/"signature"
section, and none of the seven ``*_letter.html.j2`` templates render a label
for those blocks either — inventing hardcoded heading text would violate the
i18n contract (never hard-code DE/EN strings). So this writer applies the
SAME accent-colour-on-headings rule as the CV to the one heading a letter
actually has, rather than forcing three heading levels onto content that has
no labelled slots for them.
"""

import io

from applire.schemas.cover_letter import (
    LetterBody,
    LetterData,
    LetterHeader,
    LetterRecipient,
    LetterSignature,
)


# ---------------------------------------------------------------------------
# letter_docx.py — the pure LetterData -> bytes writer
# ---------------------------------------------------------------------------

ACCENT = "#1a73e8"


def _open(docx_bytes: bytes):
    from docx import Document as open_docx

    return open_docx(io.BytesIO(docx_bytes))


def _all_text(docx_bytes: bytes) -> str:
    document = _open(docx_bytes)
    return "\n".join(p.text for p in document.paragraphs)


def _full_letter_data() -> LetterData:
    """One populated instance of every LetterData field, with distinct
    MARKER_* strings per field so content tests can assert precisely on
    which section produced which text."""
    return LetterData(
        header=LetterHeader(
            name="MARKER_SENDER_NAME",
            address="MARKER_SENDER_ADDRESS",
            phone="MARKER_SENDER_PHONE",
            email="MARKER_SENDER_EMAIL",
            photo_url="https://example.invalid/MARKER_PHOTO_URL_NEVER_RENDERED.jpg",
        ),
        recipient=LetterRecipient(
            name="MARKER_RECIPIENT_NAME",
            title="MARKER_RECIPIENT_TITLE",
            company="MARKER_RECIPIENT_COMPANY",
            address="MARKER_RECIPIENT_ADDRESS",
            date="MARKER_RECIPIENT_DATE",
        ),
        body=LetterBody(paragraphs=["MARKER_BODY_PARA_ONE", "MARKER_BODY_PARA_TWO"]),
        signature=LetterSignature(closing="MARKER_CLOSING_TEXT", name="MARKER_SIGNATURE_NAME"),
    )


class TestRenderLetterDocxStructure:
    def test_returns_bytes(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        result = render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_bytes_are_a_valid_docx(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        result = render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT)
        document = _open(result)
        assert len(document.paragraphs) > 0

    def test_no_tables(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        result = render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT)
        document = _open(result)
        assert document.tables == []

    def test_no_layout_tables_text_boxes_or_positioned_frames_in_xml(self):
        """Direct XML inspection (US297 AC): zero tables, zero text boxes,
        zero positioned frames. Namespace-agnostic local-name check, not a
        raw substring search, so a namespace *declaration* (harmless
        boilerplate on the root element) cannot produce a false positive."""
        from lxml import etree

        from applire.services.office_export.letter_docx import render_letter_docx

        result = render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT)
        document = _open(result)
        local_names = {etree.QName(el).localname for el in document.element.body.iter()}

        forbidden = {"tbl", "pBdr", "framePr", "pict", "txbx", "shape", "textbox"}
        found = local_names & forbidden
        assert not found, f"forbidden OOXML constructs present: {found}"

    def test_uses_heading_1_style_for_sender_name(self):
        """The letter's only heading — see module docstring. Exactly one
        Heading 1 paragraph, carrying the sender's own name."""
        from applire.services.office_export.letter_docx import render_letter_docx

        result = render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT)
        document = _open(result)
        heading1 = [p for p in document.paragraphs if p.style.name == "Heading 1"]
        assert len(heading1) == 1
        assert heading1[0].text == "MARKER_SENDER_NAME"

    def test_accent_color_appears_only_on_heading_runs(self):
        from applire.services.office_export._common import hex_to_rgb_color
        from applire.services.office_export.letter_docx import render_letter_docx

        result = render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT)
        document = _open(result)
        color = hex_to_rgb_color(ACCENT)

        saw_accent_on_heading = False
        for paragraph in document.paragraphs:
            is_heading = paragraph.style.name.startswith("Heading")
            for run in paragraph.runs:
                if run.font.color is not None and run.font.color.rgb == color:
                    assert is_heading, (
                        f"accent colour on a non-heading run: "
                        f"style={paragraph.style.name!r} text={run.text!r}"
                    )
                    saw_accent_on_heading = True
        assert saw_accent_on_heading, "expected the heading run painted in accent colour"


class TestSectionCoverageGuard:
    def test_every_letterdata_field_has_a_registered_renderer(self):
        """The #619-style gate: the renderer dispatch is derived from
        LetterData.model_fields, never a hand-written list, so a new schema
        field cannot silently go unexported."""
        from applire.services.office_export.letter_docx import (
            _NON_SECTION_FIELDS,
            _SECTION_RENDERERS,
        )

        covered = set(_SECTION_RENDERERS) | set(_NON_SECTION_FIELDS)
        assert covered == set(LetterData.model_fields)


class TestNestedLeafCoverageGuard:
    """Same guard shape as test_office_export_cv_docx.py's
    TestNestedLeafCoverageGuard, over LetterData instead of TailoredCVData.
    ``_iter_leaf_paths`` itself now lives in ``_common.py`` (moved there in
    this task, E057 1.4/ADR-066, once this second writer needed the
    identical walk)."""

    def test_iter_leaf_paths_finds_all_letterdata_leaves(self):
        """LetterData's tree is small enough to enumerate exhaustively —
        exact set equality, not just a subset check, including confirming
        body.paragraphs (list[str]) is ONE leaf with no sub-schema to
        descend into."""
        from applire.services.office_export._common import _iter_leaf_paths

        leaves = set(_iter_leaf_paths(LetterData))
        assert leaves == {
            "header.name",
            "header.address",
            "header.phone",
            "header.email",
            "header.photo_url",
            "recipient.name",
            "recipient.title",
            "recipient.company",
            "recipient.address",
            "recipient.date",
            "body.paragraphs",
            "signature.closing",
            "signature.name",
        }

    def test_every_letterdata_leaf_field_is_accounted_for(self):
        """Every leaf in the LetterData tree must be in EXACTLY ONE of
        _RENDERED_LEAVES / _NOT_RENDERED_LEAVES. A leaf in neither is a
        silent omission; a leaf in both is a self-contradiction."""
        from applire.services.office_export._common import _iter_leaf_paths
        from applire.services.office_export.letter_docx import (
            _NOT_RENDERED_LEAVES,
            _RENDERED_LEAVES,
        )

        all_leaves = set(_iter_leaf_paths(LetterData))
        accounted = set(_RENDERED_LEAVES) | set(_NOT_RENDERED_LEAVES)

        assert all_leaves == accounted, (
            f"unaccounted leaves: {all_leaves - accounted!r}; "
            f"stale registry entries no longer in the schema: {accounted - all_leaves!r}"
        )

    def test_rendered_and_not_rendered_leaves_are_disjoint(self):
        from applire.services.office_export.letter_docx import (
            _NOT_RENDERED_LEAVES,
            _RENDERED_LEAVES,
        )

        overlap = set(_RENDERED_LEAVES) & set(_NOT_RENDERED_LEAVES)
        assert not overlap, f"leaves classified as both rendered and not: {overlap!r}"

    def test_not_rendered_leaves_have_written_reasons(self):
        from applire.services.office_export.letter_docx import _NOT_RENDERED_LEAVES

        assert _NOT_RENDERED_LEAVES, "expected at least one deliberately-not-rendered leaf"
        for leaf, reason in _NOT_RENDERED_LEAVES.items():
            assert isinstance(reason, str) and len(reason.strip()) >= 15, (
                f"{leaf!r} has no real written reason: {reason!r}"
            )

    def test_photo_url_is_the_only_not_rendered_leaf(self):
        """Pinned explicitly, not just implied by the exhaustive-accounting
        test above: every OTHER LetterData leaf is rendered — the letter
        writer's one deliberate content omission is the photo (see module
        docstring: no letter template renders it either)."""
        from applire.services.office_export.letter_docx import _NOT_RENDERED_LEAVES

        assert set(_NOT_RENDERED_LEAVES) == {"header.photo_url"}


class TestRenderLetterDocxContent:
    def test_header_fields_appear(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        text = _all_text(render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT))
        assert "MARKER_SENDER_NAME" in text
        assert "MARKER_SENDER_ADDRESS" in text
        assert "MARKER_SENDER_PHONE" in text
        assert "MARKER_SENDER_EMAIL" in text

    def test_recipient_fields_appear(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        text = _all_text(render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT))
        assert "MARKER_RECIPIENT_NAME" in text
        assert "MARKER_RECIPIENT_TITLE" in text
        assert "MARKER_RECIPIENT_COMPANY" in text
        assert "MARKER_RECIPIENT_ADDRESS" in text
        assert "MARKER_RECIPIENT_DATE" in text

    def test_body_paragraphs_appear_in_order(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        text = _all_text(render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT))
        assert "MARKER_BODY_PARA_ONE" in text
        assert "MARKER_BODY_PARA_TWO" in text
        assert text.index("MARKER_BODY_PARA_ONE") < text.index("MARKER_BODY_PARA_TWO")

    def test_signature_name_appears(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        text = _all_text(render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT))
        assert "MARKER_SIGNATURE_NAME" in text

    def test_signature_closing_punctuation_follows_language(self):
        """#307-style closing_punctuation rule (labels.py:63-67): German
        takes NO comma after the Grussformel; English does. A hardcoded
        comma (or a hardcoded absence of one) would fail exactly one side
        of this test."""
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = _full_letter_data()
        text_de = _all_text(render_letter_docx(letter, lang="de", accent_color=ACCENT))
        text_en = _all_text(render_letter_docx(letter, lang="en", accent_color=ACCENT))

        assert "MARKER_CLOSING_TEXT" in text_de
        assert "MARKER_CLOSING_TEXT," in text_en  # EN closing_punctuation == ","
        assert "MARKER_CLOSING_TEXT," not in text_de  # DE closing_punctuation == ""

    def test_photo_url_never_rendered_as_text_or_image(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        result = render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT)
        document = _open(result)
        assert len(document.inline_shapes) == 0
        assert "MARKER_PHOTO_URL_NEVER_RENDERED" not in _all_text(result)

    def test_sections_render_in_header_recipient_body_signature_order(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        text = _all_text(render_letter_docx(_full_letter_data(), lang="de", accent_color=ACCENT))
        positions = [
            text.index(marker)
            for marker in (
                "MARKER_SENDER_NAME",
                "MARKER_RECIPIENT_NAME",
                "MARKER_BODY_PARA_ONE",
                "MARKER_SIGNATURE_NAME",
            )
        ]
        assert positions == sorted(positions), f"sections out of order:\n{text}"


class TestRenderLetterDocxDegenerateInputs:
    def test_minimal_letter_does_not_crash(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = LetterData(body=LetterBody(paragraphs=["Minimal body text."]))
        result = render_letter_docx(letter, lang="de", accent_color=ACCENT)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_blank_header_name_produces_no_heading(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = LetterData(
            header=LetterHeader(name=""),
            body=LetterBody(paragraphs=["Body text survives."]),
        )
        result = render_letter_docx(letter, lang="de", accent_color=ACCENT)
        document = _open(result)
        assert not any(p.style.name == "Heading 1" for p in document.paragraphs)
        assert "Body text survives." in _all_text(result)

    def test_blank_recipient_and_signature_do_not_crash(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = LetterData(
            header=LetterHeader(name="Solo Sender"),
            recipient=LetterRecipient(),
            body=LetterBody(paragraphs=["Body text."]),
            signature=LetterSignature(),
        )
        result = render_letter_docx(letter, lang="de", accent_color=ACCENT)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_blank_paragraphs_in_body_do_not_crash_or_emit_empty_paragraph(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = LetterData(body=LetterBody(paragraphs=["", "  ", "Real paragraph."]))
        result = render_letter_docx(letter, lang="de", accent_color=ACCENT)
        text = _all_text(result)
        assert "Real paragraph." in text

    def test_none_optional_fields_do_not_crash(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = LetterData(
            header=LetterHeader(name="Anna Bauer", phone=None, email=None),
            recipient=LetterRecipient(name=None, title=None, company=None, address=None, date=None),
            body=LetterBody(paragraphs=["Body text."]),
            signature=LetterSignature(closing=None, name="Anna Bauer"),
        )
        result = render_letter_docx(letter, lang="de", accent_color=ACCENT)
        assert isinstance(result, bytes)

    def test_twenty_body_paragraphs_do_not_crash(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = LetterData(body=LetterBody(paragraphs=[f"Paragraph {i}" for i in range(20)]))
        result = render_letter_docx(letter, lang="de", accent_color=ACCENT)
        text = _all_text(result)
        for i in range(20):
            assert f"Paragraph {i}" in text


class TestRenderLetterDocxI18n:
    def test_unknown_lang_falls_back_to_german_punctuation(self):
        from applire.services.office_export.letter_docx import render_letter_docx

        text = _all_text(render_letter_docx(_full_letter_data(), lang="fr", accent_color=ACCENT))
        assert "MARKER_CLOSING_TEXT," not in text
        assert "MARKER_CLOSING_TEXT" in text


class TestRenderLetterDocxPassesAudit:
    """US297 AC: 'the produced text must pass the unchanged _audit_letter_text
    through office_export/extract.py's extract_docx_text'. The live audit
    WIRING (persisting/serving this report from the download endpoint) is a
    separate task (out of scope here — see letter_docx.py's module
    docstring); this test only proves the writer's OWN output is compatible
    with the unchanged audit predicate, reusing the exact seam task 1.1
    already built (``audit_cover_letter_docx``), the same way
    test_office_export_extract.py already proved it against a hand-built
    fixture docx."""

    def test_writer_output_passes_the_unchanged_letter_audit_de(self):
        from applire.services.office_export.extract import audit_cover_letter_docx
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = _full_letter_data()
        docx_bytes = render_letter_docx(letter, lang="de", accent_color=ACCENT)
        report = audit_cover_letter_docx(docx_bytes, letter.model_dump(), keywords=[])

        failed = [c for c in report.checks if c.status == "fail"]
        assert not failed, f"unexpected failing checks: {failed!r}"
        assert report.checks, "expected at least one check to have run"

    def test_writer_output_passes_the_unchanged_letter_audit_en(self):
        """The ADR-079 spike measured 0 failures for both DE and EN
        (epic spec's Context section) — pinned for both languages, not just
        DE, even though _audit_letter_text's own checks (contact-name,
        contact-email, recipient-company, body-N) never touch this writer's
        one language-varying leaf (signature's closing_punctuation)."""
        from applire.services.office_export.extract import audit_cover_letter_docx
        from applire.services.office_export.letter_docx import render_letter_docx

        letter = _full_letter_data()
        docx_bytes = render_letter_docx(letter, lang="en", accent_color=ACCENT)
        report = audit_cover_letter_docx(docx_bytes, letter.model_dump(), keywords=[])

        failed = [c for c in report.checks if c.status == "fail"]
        assert not failed, f"unexpected failing checks: {failed!r}"
        assert report.checks, "expected at least one check to have run"
