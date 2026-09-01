# Copyright (C) 2024-2026 Tobias Rosenbaum
#
# This file is part of Applire.
#
# Applire is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# Applire is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with Applire. If not, see <https://www.gnu.org/licenses/>.

"""E057 Task 1.1 (US296, #637, ADR-079): the .docx text extractor + audit seam.

Synthetic fixtures only, built programmatically with python-docx (pinned 1.1.2,
requirements.txt) — no real candidate data (Documents/testdata/RealProfiles/ is
off-limits). Text boxes / mc:AlternateContent have no public python-docx API, so
those two fixture builders inject minimal hand-written OOXML via
``docx.oxml.parse_xml`` — the documented python-docx workaround pattern for
elements the library doesn't model.
"""

from io import BytesIO

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from applire.schemas.cv import TailoredCVData
from applire.services.ats_audit import _find, _norm
from applire.services.office_export.extract import (
    audit_cover_letter_docx,
    audit_cv_docx,
    extract_docx_text,
)

# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------

# wps/mc/v are not in python-docx's built-in nsdecls() map (docx/oxml/ns.py) —
# declared by hand alongside nsdecls("w", "wp", "a").
_DRAWING_EXTRA_NS = (
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
    'xmlns:v="urn:schemas-microsoft-com:vml"'
)


def _docx_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _naive_paragraphs_text(data: bytes) -> str:
    """The naive reader this task exists to replace —
    ``"\\n".join(p.text for p in doc.paragraphs)``, reproduced inline.

    This was ``cv_parser.extract_text_from_docx``'s algorithm until #640, where
    it was measured losing 9 of 9 facts on a tabellarischer Lebenslauf; that
    function now delegates to ``extract_docx_text``. The inline copy stays
    because these tests document WHY the walker exists — they must keep failing
    the naive way even though no production caller does it any more."""
    doc = Document(BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _drawing_txbx_xml(text: str) -> str:
    """A minimal DrawingML text box (``wps:txbx``) fragment carrying ``text``."""
    return (
        '<w:drawing %s %s>'
        '<wp:inline distT="0" distB="0" distL="0" distR="0">'
        '<wp:extent cx="2000000" cy="500000"/>'
        '<wp:docPr id="1" name="TextBox"/>'
        '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp><wps:cNvSpPr txBox="1"/>'
        '<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="2000000" cy="500000"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></wps:spPr>'
        f'<wps:txbx><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:txbxContent></wps:txbx>'
        '<wps:bodyPr/></wps:wsp>'
        '</a:graphicData></a:graphic></wp:inline></w:drawing>'
    ) % (nsdecls("w", "wp", "a"), _DRAWING_EXTRA_NS)


def _add_text_box(paragraph, text: str) -> None:
    """Anchor a modern DrawingML text box (no AlternateContent wrapper) to a
    fresh run in ``paragraph``."""
    run = paragraph.add_run()
    run._r.append(parse_xml(_drawing_txbx_xml(text)))


def _add_alternate_content_text_box(paragraph, text: str) -> None:
    """The REAL shape Word emits for a text box: ``mc:AlternateContent`` with a
    modern DrawingML ``mc:Choice`` branch and a legacy VML ``mc:Fallback`` branch
    — both carrying the SAME text, for cross-version-reader compatibility."""
    run = paragraph.add_run()
    # Written out in full rather than composed from _drawing_txbx_xml: nesting
    # a second xmlns-decorated <w:drawing ...> tag inside mc:Choice would
    # duplicate/conflict with the outer AlternateContent's namespace decls.
    xml = (
        '<mc:AlternateContent %s %s>'
        '<mc:Choice Requires="wps">'
        '<w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0">'
        '<wp:extent cx="2000000" cy="500000"/><wp:docPr id="2" name="TextBox"/>'
        '<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp><wps:cNvSpPr txBox="1"/>'
        '<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="2000000" cy="500000"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></wps:spPr>'
        f'<wps:txbx><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:txbxContent></wps:txbx>'
        '<wps:bodyPr/></wps:wsp></a:graphicData></a:graphic></wp:inline></w:drawing>'
        '</mc:Choice>'
        '<mc:Fallback><w:pict>'
        f'<v:shape><v:textbox><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:txbxContent></v:textbox></v:shape>'
        '</w:pict></mc:Fallback>'
        '</mc:AlternateContent>'
    ) % (nsdecls("w", "wp", "a"), _DRAWING_EXTRA_NS)
    run._r.append(parse_xml(xml))


# ---------------------------------------------------------------------------
# extract_docx_text — paragraphs, runs, tabs, breaks
# ---------------------------------------------------------------------------


def test_paragraph_order_is_preserved_with_newline_join():
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.add_paragraph("Third paragraph")
    text = extract_docx_text(_docx_bytes(doc))
    assert text.index("First paragraph") < text.index("Second paragraph") < text.index("Third paragraph")
    assert "First paragraph\n" in text  # paragraph -> newline join, not a space


def test_split_run_produces_no_spurious_separator():
    """A single visible word split across three runs at arbitrary (incl. mid-
    word) boundaries with NO inserted characters must extract as the intact
    word — 'run -> no separator' (the join rule this task must justify)."""
    doc = Document()
    p = doc.add_paragraph()
    for part in ("coord", "ina", "tion across four production clusters"):
        p.add_run(part)
    text = extract_docx_text(_docx_bytes(doc))
    assert "coordination across four production clusters" in text
    assert "coord ina tion" not in text


def test_tab_character_becomes_a_space():
    doc = Document()
    p = doc.add_paragraph("Cloudwerk GmbH")
    p.add_run().add_tab()
    p.add_run("2021–2023")
    text = extract_docx_text(_docx_bytes(doc))
    assert "Cloudwerk GmbH 2021–2023" in text


def test_line_break_becomes_a_newline():
    doc = Document()
    p = doc.add_paragraph("Berlin, Germany")
    p.add_run().add_break()  # w:br — a soft line break, not a new paragraph
    p.add_run("+49 151 1234567")
    text = extract_docx_text(_docx_bytes(doc))
    assert "Berlin, Germany\n+49 151 1234567" in text


def test_empty_document_returns_without_crashing():
    doc = Document()
    text = extract_docx_text(_docx_bytes(doc))
    assert isinstance(text, str)


# ---------------------------------------------------------------------------
# extract_docx_text — table cells (the naive .paragraphs reader misses these)
# ---------------------------------------------------------------------------


def test_table_cells_extracted_in_document_order_between_surrounding_paragraphs():
    doc = Document()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cloudwerk GmbH"
    table.cell(0, 1).text = "Senior Backend Engineer"
    table.cell(1, 0).text = "DataHaus AG"
    table.cell(1, 1).text = "Software Engineer"
    doc.add_paragraph("Outro paragraph")

    text = extract_docx_text(_docx_bytes(doc))

    # document order: intro, then the table's cells row-major, then outro —
    # NOT "every paragraph, then every table" (naive .paragraphs + .tables).
    positions = [text.index(s) for s in (
        "Intro paragraph", "Cloudwerk GmbH", "Senior Backend Engineer",
        "DataHaus AG", "Software Engineer", "Outro paragraph",
    )]
    assert positions == sorted(positions)


def test_naive_paragraphs_reader_loses_table_content():
    """Grounds the epic's own claim: .paragraphs alone (the pre-#640
    extract_text_from_docx algorithm, reproduced inline) never sees table text."""
    doc = Document()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cloudwerk GmbH"
    table.cell(0, 1).text = "Senior Backend Engineer"
    doc.add_paragraph("Outro paragraph")
    data = _docx_bytes(doc)

    naive = _naive_paragraphs_text(data)
    mine = extract_docx_text(data)

    assert "Cloudwerk GmbH" not in naive
    assert "Senior Backend Engineer" not in naive
    assert "Cloudwerk GmbH" in mine
    assert "Senior Backend Engineer" in mine


# ---------------------------------------------------------------------------
# extract_docx_text — text boxes (the naive .paragraphs reader misses these too)
# ---------------------------------------------------------------------------


def test_text_box_content_is_found():
    doc = Document()
    doc.add_paragraph("Before paragraph")
    p = doc.add_paragraph()
    _add_text_box(p, "TEXTBOX CONTENT HERE")
    doc.add_paragraph("After paragraph")

    text = extract_docx_text(_docx_bytes(doc))
    assert "TEXTBOX CONTENT HERE" in text
    assert text.index("Before paragraph") < text.index("TEXTBOX CONTENT HERE") < text.index("After paragraph")


def test_naive_paragraphs_reader_loses_text_box_content():
    doc = Document()
    p = doc.add_paragraph()
    _add_text_box(p, "SIDEBAR SKILL LIST")
    data = _docx_bytes(doc)

    naive = _naive_paragraphs_text(data)
    mine = extract_docx_text(data)

    assert "SIDEBAR SKILL LIST" not in naive
    assert "SIDEBAR SKILL LIST" in mine


def test_alternate_content_text_box_does_not_duplicate_text():
    """Word wraps a text box in mc:AlternateContent (a modern DrawingML Choice
    branch + a legacy VML Fallback branch carrying the SAME text). Walking both
    branches would report every text-box word twice."""
    doc = Document()
    p = doc.add_paragraph()
    _add_alternate_content_text_box(p, "DUPTEST")

    text = extract_docx_text(_docx_bytes(doc))
    assert text.count("DUPTEST") == 1


# ---------------------------------------------------------------------------
# Measurement (a) — ADR-079 clause 4 / the re-baselined boundary note:
# is _find's #399 kerning workaround inert on a FULL-BULLET needle against
# .docx-derived text? Never measured before this task (only ever measured for
# short structured-field needles). Three angles: (1) does exact match already
# suffice for a well-formed run split, (2) does the loose fallback still
# activate when a run split accidentally inserts a spurious space, (3) can the
# loose fallback ever produce a FALSE POSITIVE on a genuinely shortened bullet
# (the #634 shape)?
# ---------------------------------------------------------------------------

_FULL_BULLET = (
    "Coordinated the rollout of Project Phoenix across three regional R&D "
    "teams without missing a single delivery milestone."
)


def test_measurement_a_exact_match_suffices_for_clean_run_split():
    # Split at four boundaries, one of them mid-word, with ZERO inserted
    # characters — the realistic shape of a Word proofing-mark run split.
    parts = [
        "Coordinated the rollout of Project Pho",
        "enix across three region",
        "al R&D teams without miss",
        "ing a single delivery milestone.",
    ]
    assert "".join(parts) == _FULL_BULLET  # sanity: the split is lossless
    doc = Document()
    p = doc.add_paragraph()
    for part in parts:
        p.add_run(part)
    haystack = _norm(extract_docx_text(_docx_bytes(doc)))

    # The plain, non-loose substring check already succeeds — _find's #399
    # loose fallback is NOT NEEDED for a well-formed run split.
    assert haystack.find(_norm(_FULL_BULLET)) >= 0
    assert _find(_FULL_BULLET, haystack) >= 0


def test_measurement_a_loose_fallback_still_activates_on_accidental_run_split_space():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Coordinated the rollout of Project Pho")
    # An accidental leading space on the second run — the kind of split a real
    # revision-save re-flow can introduce, though not exercised by python-docx
    # itself here.
    p.add_run(" enix across three regional R&D teams without missing a single delivery milestone.")
    haystack = _norm(extract_docx_text(_docx_bytes(doc)))

    # Exact match now genuinely fails (one real extra space mid-word)...
    assert haystack.find(_norm(_FULL_BULLET)) == -1
    # ...but _find's loose fallback still finds it — the workaround is
    # REACHABLE and functional on .docx-derived text, not dead code here.
    assert _find(_FULL_BULLET, haystack) >= 0


def test_measurement_a_no_false_positive_on_genuinely_dropped_span():
    """The #634 shape, reproduced directly in a .docx: a placeholder span
    vanishes entirely (not merely mis-spaced) between the data and the
    delivered text. _find must report it ABSENT, never a false match."""
    needle = "Coordination with Project Phoenix and the R&D teams delivered the migration on time."
    mutilated = "Coordination with and the R&D teams delivered the migration on time."
    doc = Document()
    doc.add_paragraph(mutilated)
    haystack = _norm(extract_docx_text(_docx_bytes(doc)))

    assert _find(needle, haystack) == -1


# ---------------------------------------------------------------------------
# Measurement (b) — does the one measured false-positive vector (an unmapped
# ligature glyph extracting as NUL, a Chromium-PDF rendering artefact) carry
# over to .docx-derived extraction?
# ---------------------------------------------------------------------------


def test_measurement_b_ligature_codepoint_extracts_literally_never_as_nul():
    """A literal Unicode ligature codepoint (U+FB01) typed into a run's text —
    the only way a ligature reaches .docx text at all, since there is no font-
    rendering/glyph-substitution step between what <w:t> stores and what this
    extractor reads (unlike the Chromium-PDF path, where the RENDERER chooses a
    ligature glyph that pypdf must reverse-map back to Unicode, and silently
    fails to NUL when the embedded font's cmap has no entry for that glyph)."""
    doc = Document()
    doc.add_paragraph("Proﬁ Engineer")  # "Profi Engineer", fi-ligature codepoint
    text = extract_docx_text(_docx_bytes(doc))

    assert "\x00" not in text
    normalized = _norm(text)
    assert "profi engineer" in normalized  # NFKC (inside _norm) expands the ligature
    assert _find("Profi Engineer", normalized) >= 0


# ---------------------------------------------------------------------------
# The audit seam: audit_cv_docx / audit_cover_letter_docx feed the UNCHANGED
# _audit_cv_text / _audit_letter_text (🔒 boundary — ADR-066/ADR-079 clause 4).
# ---------------------------------------------------------------------------

_CV = TailoredCVData.model_validate({
    "contact": {"name": "Anna Bauer", "email": "anna@example.com", "phone": "+49 151 1234567", "location": "Berlin"},
    "summary": "Backend engineer with a cloud infrastructure focus.",
    "work_history": [
        {"company": "Cloudwerk GmbH", "role": "Senior Backend Engineer", "start_date": "2021-04", "end_date": None,
         "bullets": [
             "Built resilient FastAPI services for the payments team.",
             "Led the Kubernetes migration across four production clusters.",
         ]},
    ],
    "skills": ["Python", "FastAPI", "Kubernetes"],
    "education": [{"institution": "TU Berlin", "degree": "M.Sc.", "field": "Informatik",
                    "start_date": "2014-10", "end_date": "2017-08"}],
})


def _cv_docx_bytes(*, drop_bullet: str | None = None) -> bytes:
    """A minimal but realistic CV .docx whose paragraphs carry every field of
    ``_CV`` — except ``drop_bullet``, when given, which is deliberately omitted
    (simulating a #634-shaped writer defect) so the seam's differential check
    can be exercised against a genuine loss."""
    doc = Document()
    doc.add_paragraph("Anna Bauer")
    doc.add_paragraph("anna@example.com | +49 151 1234567 | Berlin")
    doc.add_paragraph(_CV.summary)
    doc.add_paragraph("Cloudwerk GmbH — Senior Backend Engineer (04/2021 – heute)")
    for bullet in _CV.work_history[0].bullets:
        if bullet == drop_bullet:
            continue
        doc.add_paragraph(bullet, style=None)
    doc.add_paragraph("TU Berlin — M.Sc. Informatik")
    doc.add_paragraph("Skills: " + ", ".join(_CV.skills))
    return _docx_bytes(doc)


def test_audit_cv_docx_runs_the_real_structural_checks_against_docx_text():
    report = audit_cv_docx(_cv_docx_bytes(), _CV, keywords=[])
    by_id = {c.id: c for c in report.checks}

    assert by_id["contact-name"].status == "pass"
    assert by_id["contact-email"].status == "pass"
    assert by_id["contact-phone"].status == "pass"
    assert by_id["work-0"].status == "pass"
    assert by_id["education-0"].status == "pass"
    assert by_id["skills"].status == "pass"
    # content-0 = summary, content-1/2 = the two work-entry bullets (_free_text_snippets order).
    assert by_id["content-0"].status == "pass"
    assert by_id["content-1"].status == "pass"
    assert by_id["content-2"].status == "pass"
    assert report.document == "cv"


def test_audit_cv_docx_detects_a_genuinely_dropped_bullet():
    dropped = _CV.work_history[0].bullets[1]
    report = audit_cv_docx(_cv_docx_bytes(drop_bullet=dropped), _CV, keywords=[])
    by_id = {c.id: c for c in report.checks}

    assert by_id["content-2"].status == "fail"
    # Nothing else about the differential regresses — this is a targeted loss,
    # not a wholesale extraction failure.
    assert by_id["content-1"].status == "pass"
    assert by_id["contact-name"].status == "pass"


def test_audit_cv_docx_page_band_is_reported_not_applicable():
    """ADR-079 clause 4: the band is present and explicitly inapplicable.

    This test was written as a CHARACTERIZATION test pinning the opposite —
    that the band was simply ABSENT — because the decision between amending the
    audit functions and constructing the check in this seam was still open. Its
    own docstring said: *"If this test ever goes red because someone added a
    page-length entry, that is progress — update it deliberately, don't just
    delete it."* The founder chose the amendment on 2026-09-01, the test went
    red on the genuine fix, and this is that deliberate update. The pin is kept
    rather than dropped because the property it guards has not gone away — it
    has inverted: an absent band is invisible to `passed` and `failed` alike,
    so a report would read clean on something never evaluated (the #634 shape),
    and that must stay impossible."""
    report = audit_cv_docx(_cv_docx_bytes(), _CV, keywords=[])

    band = next((c for c in report.checks if c.id == "page-length"), None)
    assert band is not None, "the band must never be silently omitted (ADR-079 cl. 4)"
    assert band.status == "not_applicable"
    assert band.details_key, "the reason must be machine-readable for the frontend"

    # It must land in its own bucket, not in either numerator.
    assert report.passed == sum(1 for c in report.checks if c.status == "pass")
    assert report.failed == sum(1 for c in report.checks if c.status == "fail")
    assert report.not_applicable == 1
    assert report.passed + report.failed + report.not_applicable == len(report.checks)


_LETTER = {
    "header": {"name": "Anna Bauer", "email": "anna@example.com", "phone": None, "address": "Berlin"},
    "recipient": {"company": "Cloudwerk GmbH", "name": "Herr Schmidt", "title": None, "address": None, "date": None},
    "body": {"paragraphs": [
        "Sehr geehrter Herr Schmidt,",
        "mit grossem Interesse habe ich Ihre Stellenanzeige fuer die Position als Senior Backend Engineer gelesen.",
    ]},
    "signature": {"name": "Anna Bauer"},
}


def _letter_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Anna Bauer")
    doc.add_paragraph("anna@example.com")
    doc.add_paragraph("Cloudwerk GmbH")
    for para in _LETTER["body"]["paragraphs"]:
        doc.add_paragraph(para)
    doc.add_paragraph("Anna Bauer")
    return _docx_bytes(doc)


def test_audit_cover_letter_docx_runs_the_real_structural_checks_against_docx_text():
    report = audit_cover_letter_docx(_letter_docx_bytes(), _LETTER, keywords=[])
    by_id = {c.id: c for c in report.checks}

    assert by_id["contact-name"].status == "pass"
    assert by_id["contact-email"].status == "pass"
    assert by_id["recipient-company"].status == "pass"
    assert by_id["body-0"].status == "pass"
    assert by_id["body-1"].status == "pass"
    assert report.document == "cover_letter"
    # ADR-079 cl. 4, letter side: present and explicitly inapplicable, never absent.
    band = next((c for c in report.checks if c.id == "page-length"), None)
    assert band is not None and band.status == "not_applicable"
    assert report.not_applicable == 1
